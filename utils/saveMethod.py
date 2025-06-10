"""
Word Document Save Methods

This module provides functions for saving Word documents in different formats.
"""
import os
import sys

# 标记库是否已安装
docx_installed = True

# 尝试导入python-docx库，如果没有安装则标记为未安装但不退出
try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENTATION, WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("警告: 未检测到python-docx库，Word文档功能将不可用")
    print("请使用以下命令安装: pip install python-docx")
    docx_installed = False


def save_document_as_pdf(file_path: str) -> str:
    """
    将Word文档保存为PDF格式。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法导出PDF，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 构建PDF文件路径
        pdf_path = os.path.splitext(file_path)[0] + ".pdf"
        
        # 尝试使用Microsoft Word COM对象导出PDF
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17表示PDF格式
            doc.Close()
            word.Quit()
            
            return f"成功将文档导出为PDF: {os.path.basename(pdf_path)}"
        
        except ImportError:
            # 如果没有win32com库或不是Windows系统，返回错误信息
            return "错误: 导出PDF功能需要在Windows系统上安装pywin32库，请使用以下命令安装: pip install pywin32"
    
    except Exception as e:
        return f"导出PDF时出错: {str(e)}"


def save_document_as(file_path: str, output_format: str = "docx", new_filename: str = None) -> str:
    """
    将Word文档保存为指定格式。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        output_format: 输出格式，可选值: "docx", "doc", "pdf", "txt", "html"
        new_filename: 新文件名(不含扩展名)，如果不提供则使用原文件名
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法保存文档，请先安装python-docx库: pip install python-docx"
    
    # 检查格式是否支持
    supported_formats = ["docx", "doc", "pdf", "txt", "html"]
    if output_format.lower() not in supported_formats:
        return f"错误: 不支持的输出格式 '{output_format}'，可选值为: {', '.join(supported_formats)}"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 构建新文件路径
        original_basename = os.path.splitext(os.path.basename(file_path))[0]
        output_dirname = os.path.dirname(file_path)
        
        # 如果提供了新文件名，则使用新文件名
        if new_filename:
            output_basename = new_filename
        else:
            output_basename = original_basename
        
        # 创建新文件的完整路径
        output_path = os.path.join(output_dirname, f"{output_basename}.{output_format}")
        
        # 根据输出格式选择不同的处理方式
        if output_format.lower() == "pdf":
            # 使用前面定义的导出PDF功能
            return save_document_as_pdf(file_path)
        
        elif output_format.lower() in ["docx", "doc"]:
            try:
                # 尝试使用Microsoft Word COM对象保存
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # 使用数字格式指定不同的Word格式
                format_map = {
                    "docx": 16,  # wdFormatDocumentDefault (*.docx)
                    "doc": 0     # wdFormatDocument97 (*.doc)
                }
                
                doc = word.Documents.Open(file_path)
                doc.SaveAs(output_path, FileFormat=format_map[output_format.lower()])
                doc.Close()
                word.Quit()
                
                return f"成功将文档保存为 {output_format} 格式: {os.path.basename(output_path)}"
            
            except ImportError:
                # 如果是docx格式，我们可以使用python-docx直接保存
                if output_format.lower() == "docx":
                    doc = Document(file_path)
                    doc.save(output_path)
                    return f"成功将文档保存为 DOCX 格式: {os.path.basename(output_path)}"
                else:
                    return "错误: 保存为DOC格式需要在Windows系统上安装pywin32库"
        
        elif output_format.lower() == "txt":
            # 将文档转换为纯文本
            doc = Document(file_path)
            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return f"成功将文档保存为文本格式: {os.path.basename(output_path)}"
        
        elif output_format.lower() == "html":
            try:
                # 尝试使用Microsoft Word COM对象保存为HTML
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(file_path)
                doc.SaveAs(output_path, FileFormat=8)  # 8表示HTML格式
                doc.Close()
                word.Quit()
                
                return f"成功将文档保存为HTML格式: {os.path.basename(output_path)}"
            
            except ImportError:
                return "错误: 保存为HTML格式需要在Windows系统上安装pywin32库"
    
    except Exception as e:
        return f"保存文档时出错: {str(e)}"


