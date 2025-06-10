"""
媒体和表格操作模块 - 集成图片插入和表格操作功能
支持批量操作以提高效率
"""

import os
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def batch_insert_images(
    file_path: str,
    images_data: List[Dict[str, Any]],
    output_path: Optional[str] = None
) -> str:
    """
    批量插入图片到Word文档
    
    Args:
        file_path: Word文档路径
        images_data: 图片数据列表，每个元素包含：
            - image_path: 图片文件路径（必需）
            - width: 图片宽度厘米（可选）
            - height: 图片高度厘米（可选）
            - after_paragraph: 插入位置段落索引（可选，默认-1表示末尾）
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        doc = Document(file_path)
        success_count = 0
        failed_operations = []
        
        # 批量处理图片插入
        for i, img_data in enumerate(images_data):
            try:
                image_path = img_data.get('image_path', '')
                width = img_data.get('width')
                height = img_data.get('height')
                after_paragraph = img_data.get('after_paragraph', -1)
                
                # 处理图片路径
                if not os.path.isabs(image_path):
                    image_path = os.path.join(output_path, image_path)
                
                # 检查图片文件是否存在
                if not os.path.exists(image_path):
                    failed_operations.append((i, f"图片文件 {image_path} 不存在"))
                    continue
                
                # 确定插入位置
                if after_paragraph == -1:
                    paragraph = doc.add_paragraph()
                elif 0 <= after_paragraph < len(doc.paragraphs):
                    target_paragraph = doc.paragraphs[after_paragraph]
                    paragraph = doc.add_paragraph()
                    # 移动段落到指定位置
                    target_paragraph._p.addnext(paragraph._p)
                else:
                    failed_operations.append((i, f"无效的段落索引: {after_paragraph}"))
                    continue
                
                # 插入图片
                run = paragraph.add_run()
                if width and height:
                    run.add_picture(image_path, width=Cm(width), height=Cm(height))
                elif width:
                    run.add_picture(image_path, width=Cm(width))
                elif height:
                    run.add_picture(image_path, height=Cm(height))
                else:
                    run.add_picture(image_path)
                
                success_count += 1
                
            except Exception as e:
                failed_operations.append((i, str(e)))
        
        # 保存文档
        doc.save(file_path)
        
        result_msg = f"成功批量插入 {success_count} 张图片到文档 {os.path.basename(file_path)}"
        if failed_operations:
            result_msg += f"，但有 {len(failed_operations)} 个操作失败"
        
        return result_msg
        
    except Exception as e:
        return f"批量插入图片时出错: {str(e)}"

def batch_insert_tables(
    file_path: str,
    tables_data: List[Dict[str, Any]],
    output_path: Optional[str] = None
) -> str:
    """
    批量插入表格到Word文档
    
    Args:
        file_path: Word文档路径
        tables_data: 表格数据列表，每个元素包含：
            - rows: 表格行数（必需）
            - cols: 表格列数（必需）
            - data: 表格内容二维数组（可选）
            - after_paragraph: 插入位置段落索引（可选，默认-1表示末尾）
            - style: 表格样式（可选，默认"Table Grid"）
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        doc = Document(file_path)
        success_count = 0
        failed_operations = []
        
        # 批量处理表格插入
        for i, table_data in enumerate(tables_data):
            try:
                rows = table_data.get('rows', 0)
                cols = table_data.get('cols', 0)
                data = table_data.get('data')
                after_paragraph = table_data.get('after_paragraph', -1)
                style = table_data.get('style', 'Table Grid')
                
                # 验证表格尺寸
                if rows <= 0 or cols <= 0:
                    failed_operations.append((i, "表格行数和列数必须大于0"))
                    continue
                
                # 验证段落索引
                if after_paragraph != -1 and (after_paragraph < 0 or after_paragraph >= len(doc.paragraphs)):
                    failed_operations.append((i, f"无效的段落索引: {after_paragraph}"))
                    continue
                
                # 创建表格
                if after_paragraph == -1:
                    table = doc.add_table(rows=rows, cols=cols)
                else:
                    target_paragraph = doc.paragraphs[after_paragraph]
                    table = doc.add_table(rows=rows, cols=cols)
                    # 移动表格到指定位置
                    target_paragraph._p.addnext(table._tbl)
                
                # 设置表格样式
                table.style = style
                
                # 填充表格数据
                if data:
                    for row_idx, row_data in enumerate(data):
                        if row_idx < rows:
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < cols:
                                    table.cell(row_idx, col_idx).text = str(cell_data)
                
                success_count += 1
                
            except Exception as e:
                failed_operations.append((i, str(e)))
        
        # 保存文档
        doc.save(file_path)
        
        result_msg = f"成功批量插入 {success_count} 个表格到文档 {os.path.basename(file_path)}"
        if failed_operations:
            result_msg += f"，但有 {len(failed_operations)} 个操作失败"
        
        return result_msg
        
    except Exception as e:
        return f"批量插入表格时出错: {str(e)}"

def batch_edit_table_cells(
    file_path: str,
    edit_operations: List[Dict[str, Any]],
    output_path: Optional[str] = None
) -> str:
    """
    批量编辑表格单元格内容
    
    Args:
        file_path: Word文档路径
        edit_operations: 编辑操作列表，每个元素包含：
            - table_index: 表格索引（必需，从0开始）
            - cell_edits: 单元格编辑列表，每个元素包含：
                - row: 行索引（从0开始）
                - col: 列索引（从0开始）
                - text: 单元格内容
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        doc = Document(file_path)
        total_edited = 0
        failed_operations = []
        
        # 批量处理表格编辑操作
        for i, operation in enumerate(edit_operations):
            try:
                table_index = operation.get('table_index', -1)
                cell_edits = operation.get('cell_edits', [])
                
                # 验证表格索引
                if table_index < 0 or table_index >= len(doc.tables):
                    failed_operations.append((i, f"无效的表格索引: {table_index}"))
                    continue
                
                table = doc.tables[table_index]
                
                # 处理单元格编辑
                for cell_edit in cell_edits:
                    row = cell_edit.get('row', -1)
                    col = cell_edit.get('col', -1)
                    text = cell_edit.get('text', '')
                    
                    # 验证行列索引
                    if row < 0 or row >= len(table.rows):
                        failed_operations.append((i, f"无效的行索引: {row}"))
                        continue
                    
                    if col < 0 or col >= len(table.columns):
                        failed_operations.append((i, f"无效的列索引: {col}"))
                        continue
                    
                    # 编辑单元格
                    table.cell(row, col).text = text
                    total_edited += 1
                
            except Exception as e:
                failed_operations.append((i, str(e)))
        
        # 保存文档
        doc.save(file_path)
        
        result_msg = f"成功批量编辑 {total_edited} 个表格单元格"
        if failed_operations:
            result_msg += f"，但有 {len(failed_operations)} 个操作失败"
        
        return result_msg
        
    except Exception as e:
        return f"批量编辑表格单元格时出错: {str(e)}"
def insert_table_of_contents(
    file_path: str,
    title: str = "目录",
    levels: int = 3,
    after_paragraph: int = 0,
    output_path: Optional[str] = None
) -> str:
    """
    在Word文档中插入目录。
    
    Args:
        file_path: Word文档路径
        title: 目录标题
        levels: 目录级别数 (1-9)
        after_paragraph: 在指定段落后插入目录，默认为文档开头第一段后
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 校验参数
    if levels < 1 or levels > 9:
        return "错误: 目录级别数必须在1至9之间"
    
    try:
        # 尝试使用Word COM对象添加目录
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            
            # 将光标移动到指定段落后
            if after_paragraph >= 0 and after_paragraph < doc.Paragraphs.Count:
                range_to_insert = doc.Paragraphs(after_paragraph + 1).Range
                range_to_insert.Collapse(0)
                
                # 插入换行符创建新段落
                range_to_insert.InsertParagraphAfter()
                range_to_insert.Collapse(0)
                
                # 插入标题
                if title:
                    range_to_insert.Text = title
                    range_to_insert.InsertParagraphAfter()
                    range_to_insert.Collapse(0)
                
                # 插入目录
                toc_range = range_to_insert
                toc_range.Fields.Add(Range=toc_range, Type=-1, Text=f"TOC \\o \"1-{levels}\" \\h", PreserveFormatting=True)
            else:
                # 在文档开头插入目录
                range_to_insert = doc.Paragraphs(1).Range
                range_to_insert.Collapse(1)
                
                # 插入标题
                if title:
                    range_to_insert.Text = title
                    range_to_insert.InsertParagraphAfter()
                    range_to_insert.Collapse(0)
                
                # 插入目录
                toc_range = range_to_insert
                toc_range.Fields.Add(Range=toc_range, Type=-1, Text=f"TOC \\o \"1-{levels}\" \\h", PreserveFormatting=True)
            
            # 更新目录
            if doc.TablesOfContents.Count > 0:
                doc.TablesOfContents(1).Update()
            
            # 保存并关闭
            doc.Save()
            doc.Close()
            word.Quit()
            
            pythoncom.CoUninitialize()
            
            return f"成功在文档 {os.path.basename(file_path)} 中插入目录"
        
        except ImportError:
            # 使用python-docx的方式添加目录（功能受限）
            doc = Document(file_path)
            
            # 检查指定段落是否有效
            if after_paragraph >= len(doc.paragraphs):
                return f"错误: 无效的段落索引 {after_paragraph}，文档共有 {len(doc.paragraphs)} 个段落"
            
            # 在指定位置插入目录标题
            if after_paragraph == 0:
                # 在文档开头插入
                if title:
                    heading_para = doc.add_paragraph(title, style="Heading 1")
                    if len(doc.paragraphs) > 1:
                        first_para = doc.paragraphs[1]._p
                        heading_para._p.addnext(first_para)
            else:
                # 在指定段落后插入
                paragraph = doc.paragraphs[after_paragraph]
                if title:
                    new_para = doc.add_paragraph()
                    paragraph._p.addnext(new_para._p)
                    new_para.text = title
                    new_para.style = "Heading 1"
            
            # 创建目录字段
            toc_para = doc.add_paragraph()
            if title:
                if after_paragraph == 0 and len(doc.paragraphs) > 1:
                    doc.paragraphs[1]._p.addnext(toc_para._p)
                elif after_paragraph > 0:
                    doc.paragraphs[after_paragraph + 1]._p.addnext(toc_para._p)
            else:
                if after_paragraph < len(doc.paragraphs):
                    doc.paragraphs[after_paragraph]._p.addnext(toc_para._p)
            
            toc_run = toc_para.add_run()
            
            # 添加目录字段XML
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            toc_run._r.append(fldChar)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = f' TOC \\o "1-{levels}" \\h '
            toc_run._r.append(instrText)
            
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            toc_run._r.append(fldChar)
            
            # 保存文档
            doc.save(file_path)
            
            return f"成功在文档 {os.path.basename(file_path)} 中插入目录（需要在Word中手动更新）"
    
    except Exception as e:
        return f"插入目录时出错: {str(e)}"
