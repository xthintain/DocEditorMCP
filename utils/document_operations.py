"""
Word Document Operations

This module provides functions for opening, reading, and closing Word documents.
"""
import os
import sys

# 标记库是否已安装
docx_installed = True

# 尝试导入python-docx库，如果没有安装则标记为未安装但不退出
try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENTATION, WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("警告: 未检测到python-docx库，Word文档功能将不可用")
    print("请使用以下命令安装: pip install python-docx")
    docx_installed = False


def open_and_read_word_document(file_path: str) -> str:
    """
    打开并读取Word文档的完整内容。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
    
    Returns:
        文档的完整内容
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法读取Word文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 提取文档基本信息
        paragraphs = [p.text for p in doc.paragraphs]
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        # 构建文档信息头
        doc_info = (
            f"文件名: {os.path.basename(file_path)}\n"
            f"段落数: {len(paragraphs)}\n"
            f"标题数: {len(headings)}\n\n"
        )
        
        # 构建完整文档内容，保留段落结构，并在每段前添加段落编号
        full_content = ""
        for i, p_text in enumerate(paragraphs):
            # 添加段落编号 (i) 在每段前
            full_content += f"[{i}] {p_text}\n"
        
        # 返回文档信息和完整内容
        return doc_info + full_content
    except Exception as e:
        return f"读取Word文档时出错: {str(e)}"


def close_document(file_path: str, save_changes: bool = True) -> str:
    """
    关闭Word文档，可选是否保存更改。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        save_changes: 是否保存更改，默认为True
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法关闭文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 尝试使用Microsoft Word COM对象关闭文档
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            
            # 检查文档是否已经打开
            doc_found = False
            for doc in word.Documents:
                if os.path.abspath(doc.FullName) == os.path.abspath(file_path):
                    if save_changes:
                        doc.Save()
                    doc.Close(SaveChanges=save_changes)
                    doc_found = True
                    break
            
            if not doc_found:
                # 如果文档未打开，则python-docx中没有明确的关闭方法
                # 这里我们模拟一个"关闭"操作，即打开文档后立即关闭
                doc = Document(file_path)
                if save_changes:
                    doc.save(file_path)
            
            pythoncom.CoUninitialize()
            
            return f"成功关闭文档: {os.path.basename(file_path)}" + (" 并保存更改" if save_changes else "")
        
        except ImportError:
            # 如果没有win32com库，使用python-docx的方式
            doc = Document(file_path)
            if save_changes:
                doc.save(file_path)
            
            # python-docx没有明确的关闭方法，垃圾收集器会处理
            return f"成功关闭文档: {os.path.basename(file_path)}" + (" 并保存更改" if save_changes else "")
    
    except Exception as e:
        return f"关闭文档时出错: {str(e)}"
