"""
文档格式化操作模块 - 页眉页脚、页面布局和文档合并功能
"""

import os
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 检查docx库安装状态
try:
    import docx
    docx_installed = True
except ImportError:
    docx_installed = False

def add_header_footer(
    file_path: str,
    header_text: str = None,
    footer_text: str = None,
    page_numbers: bool = False,
    output_path: Optional[str] = None
) -> str:
    """
    为Word文档添加页眉和页脚。
    
    Args:
        file_path: Word文档路径
        header_text: 页眉文本（可选）
        footer_text: 页脚文本（可选）
        page_numbers: 是否在页脚添加页码
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法添加页眉页脚，请先安装python-docx库: pip install python-docx"
    
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 检查是否提供了有效的参数
    if header_text is None and footer_text is None and not page_numbers:
        return "错误: 请至少提供页眉文本、页脚文本或启用页码"
    
    try:
        # 尝试使用Word COM对象添加页眉页脚（功能最完整）
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            
            # 添加页眉
            if header_text:
                for section in range(1, doc.Sections.Count + 1):
                    header = doc.Sections(section).Headers(1)
                    header.Range.Text = header_text
            
            # 添加页脚
            if footer_text or page_numbers:
                for section in range(1, doc.Sections.Count + 1):
                    footer = doc.Sections(section).Footers(1)
                    
                    if footer_text:
                        footer.Range.Text = footer_text
                    
                    if page_numbers:
                        footer.PageNumbers.Add()
            
            doc.Save()
            doc.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            
            return f"成功为文档 {os.path.basename(file_path)} 添加页眉页脚"
        
        except ImportError:
            # 使用python-docx的方式添加页眉页脚（功能受限）
            doc = Document(file_path)
            
            for section in doc.sections:
                if header_text:
                    header = section.header
                    header_para = header.paragraphs[0]
                    header_para.text = header_text
                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if footer_text:
                    footer = section.footer
                    footer_para = footer.paragraphs[0]
                    footer_para.text = footer_text
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if page_numbers:
                    footer = section.footer
                    footer_para = footer.paragraphs[0] if footer_text else footer.add_paragraph()
                    
                    run = footer_para.add_run()
                    
                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(qn('w:fldCharType'), 'begin')
                    run._r.append(fldChar)
                    
                    instrText = OxmlElement('w:instrText')
                    instrText.text = ' PAGE '
                    run._r.append(instrText)
                    
                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(qn('w:fldCharType'), 'end')
                    run._r.append(fldChar)
                    
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.save(file_path)
            
            return f"成功为文档 {os.path.basename(file_path)} 添加页眉页脚"
    
    except Exception as e:
        return f"添加页眉页脚时出错: {str(e)}"

def set_page_layout(
    file_path: str,
    orientation: str = None,
    page_width: float = None,
    page_height: float = None,
    left_margin: float = None,
    right_margin: float = None,
    top_margin: float = None,
    bottom_margin: float = None,
    section_indices: List[int] = None,
    apply_to_all: bool = False,
    output_path: Optional[str] = None
) -> str:
    """
    设置Word文档的页面布局，支持单个或多个节的批量设置。
    
    Args:
        file_path: Word文档路径
        orientation: 页面方向，可选值: "portrait"(纵向), "landscape"(横向)
        page_width: 页面宽度（厘米）
        page_height: 页面高度（厘米）
        left_margin: 左边距（厘米）
        right_margin: 右边距（厘米）
        top_margin: 上边距（厘米）
        bottom_margin: 下边距（厘米）
        section_indices: 要设置的节索引列表（从0开始），如果为None且apply_to_all为False则只设置第一节
        apply_to_all: 是否应用到所有节，默认为False
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    if not docx_installed:
        return "错误: 无法设置页面布局，请先安装python-docx库: pip install python-docx"
    
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 校验方向参数
    orientation_map = {
        "portrait": WD_ORIENTATION.PORTRAIT,
        "landscape": WD_ORIENTATION.LANDSCAPE
    }
    
    if orientation and orientation.lower() not in orientation_map:
        return f"错误: 无效的页面方向 '{orientation}'，可选值为: portrait, landscape"
    
    try:
        doc = Document(file_path)
        total_sections = len(doc.sections)
        
        # 确定要处理的节索引
        if apply_to_all:
            # 应用到所有节
            target_indices = list(range(total_sections))
        elif section_indices is not None:
            # 使用指定的节索引列表
            target_indices = []
            invalid_indices = []
            
            for idx in section_indices:
                if 0 <= idx < total_sections:
                    target_indices.append(idx)
                else:
                    invalid_indices.append(idx)
            
            if invalid_indices:
                return f"错误: 无效的节索引 {invalid_indices}，文档共有 {total_sections} 节"
            
            if not target_indices:
                return "错误: 没有有效的节索引"
        else:
            # 默认只设置第一节
            target_indices = [0] if total_sections > 0 else []
        
        if not target_indices:
            return "错误: 文档中没有可设置的节"
        
        # 记录成功设置的节数量
        success_count = 0
        failed_indices = []
        
        # 批量设置页面布局
        for section_index in target_indices:
            try:
                section = doc.sections[section_index]
                
                # 设置页面方向
                if orientation:
                    section.orientation = orientation_map[orientation.lower()]
                
                # 设置页面尺寸
                if page_width and page_height:
                    section.page_width = Cm(page_width)
                    section.page_height = Cm(page_height)
                
                # 设置页边距
                if left_margin is not None:
                    section.left_margin = Cm(left_margin)
                
                if right_margin is not None:
                    section.right_margin = Cm(right_margin)
                
                if top_margin is not None:
                    section.top_margin = Cm(top_margin)
                
                if bottom_margin is not None:
                    section.bottom_margin = Cm(bottom_margin)
                
                success_count += 1
                
            except Exception as e:
                failed_indices.append((section_index, str(e)))
        
        # 保存文档
        doc.save(file_path)
        
        # 生成结果消息
        if apply_to_all:
            result_msg = f"成功设置文档 {os.path.basename(file_path)} 所有 {success_count} 节的页面布局"
        elif len(target_indices) == 1:
            result_msg = f"成功设置文档 {os.path.basename(file_path)} 第 {target_indices[0]+1} 节的页面布局"
        else:
            result_msg = f"成功设置文档 {os.path.basename(file_path)} 中 {success_count} 个节的页面布局"
        
        if failed_indices:
            result_msg += f"，但有 {len(failed_indices)} 个节设置失败"
        
        return result_msg
        
    except Exception as e:
        return f"设置页面布局时出错: {str(e)}"

def merge_documents(
    main_file_path: str,
    files_to_merge: List[str],
    output_path: Optional[str] = None
) -> str:
    """
    合并多个Word文档。
    
    Args:
        main_file_path: 主文档路径（合并后的文档将保存为该文件）
        files_to_merge: 要合并的文档路径列表
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    if not docx_installed:
        return "错误: 无法合并文档，请先安装python-docx库: pip install python-docx"
    
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(main_file_path):
        main_file_path = os.path.join(output_path, main_file_path)
    
    if not files_to_merge:
        return "错误: 请提供至少一个要合并的文档"
    
    # 处理文件路径
    processed_files = []
    for file_path in files_to_merge:
        if not os.path.isabs(file_path):
            file_path = os.path.join(output_path, file_path)
        
        if not os.path.exists(file_path):
            return f"错误: 文件 {file_path} 不存在"
        
        processed_files.append(file_path)
    
    try:
        # 尝试使用Word COM对象合并文档（功能最完整）
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            if not os.path.exists(main_file_path):
                doc = word.Documents.Add()
                doc.SaveAs(main_file_path)
            else:
                doc = word.Documents.Open(main_file_path)
            
            merged_count = 0
            
            for file_path in processed_files:
                word.Selection.EndKey(Unit=6)
                
                if merged_count > 0:
                    word.Selection.InsertBreak(Type=2)
                
                word.Selection.InsertFile(file_path)
                merged_count += 1
            
            doc.Save()
            doc.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            
            return f"成功将 {merged_count} 个文档合并到 {os.path.basename(main_file_path)}"
        
        except ImportError:
            # 使用python-docx方式合并文档（功能受限）
            if os.path.exists(main_file_path):
                main_doc = Document(main_file_path)
            else:
                main_doc = Document()
            
            merged_count = 0
            
            for file_path in processed_files:
                doc_to_merge = Document(file_path)
                
                if merged_count > 0:
                    main_doc.add_section()
                
                for paragraph in doc_to_merge.paragraphs:
                    new_paragraph = main_doc.add_paragraph()
                    
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name
                    
                    new_paragraph.style = paragraph.style
                    new_paragraph.alignment = paragraph.alignment
                
                for table in doc_to_merge.tables:
                    new_table = main_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style
                    
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = cell.text
                
                merged_count += 1
            
            main_doc.save(main_file_path)
            
            return f"成功将 {merged_count} 个文档合并到 {os.path.basename(main_file_path)}"
    
    except Exception as e:
        return f"合并文档时出错: {str(e)}"

def apply_consistent_formatting(
    file_path: str,
    content_type: str = "heading",  # "heading", "title", "normal" (正文)
    level: int = 1,                 # 对标题有效，表示几级标题
    font_name: str = None,
    font_size: float = None,
    bold: bool = None,
    italic: bool = None,
    underline: bool = None,
    font_color: str = None,
    before_spacing: float = None,
    after_spacing: float = None,
    line_spacing: float = None,
    line_spacing_rule: str = "multiple"
) -> str:
    """
    将指定的格式应用到所有同级标题或正文，确保格式一致性。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        content_type: 内容类型，可选值为"heading"（标题）, "title"（文档标题）, "normal"（正文）
        level: 标题级别（仅当content_type为"heading"时有效），1-9
        font_name: 字体名称
        font_size: 字体大小（磅值）
        bold: 是否加粗
        italic: 是否斜体
        underline: 是否下划线
        font_color: 字体颜色（十六进制格式，如"#FF0000"）
        before_spacing: 段前间距（磅值）
        after_spacing: 段后间距（磅值）
        line_spacing: 行间距值
        line_spacing_rule: 行间距规则（"multiple"/"exact"/"atLeast"）
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法应用格式，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 根据内容类型和级别应用格式
        target_style = None
        format_count = 0
        
        if content_type.lower() == "heading":
            # 为标题设置格式
            target_style = f"Heading {level}"
            for paragraph in doc.paragraphs:
                if paragraph.style.name == target_style:
                    _apply_format_to_paragraph(
                        paragraph, font_name, font_size, bold, italic, underline, 
                        font_color, before_spacing, after_spacing, line_spacing, line_spacing_rule
                    )
                    format_count += 1
                    
        elif content_type.lower() == "title":
            # 为文档标题设置格式
            target_style = "Title"
            for paragraph in doc.paragraphs:
                if paragraph.style.name == target_style:
                    _apply_format_to_paragraph(
                        paragraph, font_name, font_size, bold, italic, underline, 
                        font_color, before_spacing, after_spacing, line_spacing, line_spacing_rule
                    )
                    format_count += 1
                    
        elif content_type.lower() == "normal":
            # 为正文设置格式
            for paragraph in doc.paragraphs:
                if paragraph.style.name == "Normal" or paragraph.style.name.startswith("Body"):
                    _apply_format_to_paragraph(
                        paragraph, font_name, font_size, bold, italic, underline, 
                        font_color, before_spacing, after_spacing, line_spacing, line_spacing_rule
                    )
                    format_count += 1
                    
        else:
            return f"错误: 不支持的内容类型 '{content_type}'，请使用 'heading', 'title' 或 'normal'"
        
        # 保存文档
        doc.save(file_path)
        
        content_type_str = {
            "heading": f"{level}级标题",
            "title": "文档标题",
            "normal": "正文"
        }.get(content_type.lower(), content_type)
        
        return f"成功为 {format_count} 个{content_type_str}应用一致格式"
        
    except Exception as e:
        return f"应用一致格式时出错: {str(e)}"

def _apply_format_to_paragraph(
    paragraph, font_name, font_size, bold, italic, underline, 
    font_color, before_spacing, after_spacing, line_spacing, line_spacing_rule
):
    """
    向段落应用指定的格式。
    
    Args:
        paragraph: 目标段落对象
        font_name: 字体名称
        font_size: 字体大小
        bold: 是否加粗
        italic: 是否斜体
        underline: 是否下划线
        font_color: 字体颜色
        before_spacing: 段前间距
        after_spacing: 段后间距
        line_spacing: 行间距值
        line_spacing_rule: 行间距规则
    """
    # 应用字体样式
    if any([font_name, font_size, bold is not None, italic is not None, underline is not None, font_color]):
        for run in paragraph.runs:
            if font_name:
                run.font.name = font_name
                # 设置中文字体
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                except:
                    pass
                
            if font_size:
                run.font.size = Pt(font_size)
                
            if bold is not None:
                run.font.bold = bold
                
            if italic is not None:
                run.font.italic = italic
                
            if underline is not None:
                run.font.underline = underline
                
            if font_color:
                try:
                    # 解析十六进制颜色
                    color = font_color.lstrip('#')
                    rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                    run.font.color.rgb = RGBColor(*rgb)
                except:
                    pass
    
    # 应用段落间距
    if any([before_spacing is not None, after_spacing is not None, line_spacing is not None]):
        if before_spacing is not None:
            paragraph.paragraph_format.space_before = Pt(before_spacing)
            
        if after_spacing is not None:
            paragraph.paragraph_format.space_after = Pt(after_spacing)
            
        if line_spacing is not None:
            if line_spacing_rule.lower() == "multiple":
                paragraph.paragraph_format.line_spacing = line_spacing
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            elif line_spacing_rule.lower() == "exact":
                paragraph.paragraph_format.line_spacing = Pt(line_spacing)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            elif line_spacing_rule.lower() == "atleast":
                paragraph.paragraph_format.line_spacing = Pt(line_spacing)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
