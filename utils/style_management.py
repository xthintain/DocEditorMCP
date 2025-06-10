"""
文档样式管理模块 - 实现样式的创建、应用和管理功能
"""

import os
from typing import List, Dict, Any, Optional, Union
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants
import json

# 检查docx库安装状态
try:
    import docx
    docx_installed = True
except ImportError:
    docx_installed = False

# 尝试导入pywin32，用于某些高级功能
try:
    import win32com.client
    import pythoncom
    win32com_installed = True
except ImportError:
    win32com_installed = False

def create_custom_style(
    file_path: str,
    style_name: str,
    style_type: str = "paragraph",  # paragraph, character, table, list
    based_on: str = None,
    font_name: str = None,
    font_size: float = None,
    font_bold: bool = None,
    font_italic: bool = None,
    font_underline: bool = None,
    font_color: str = None,  # 十六进制颜色，如 "#000000"
    alignment: str = None,  # left, right, center, justify
    line_spacing: float = None,
    space_before: float = None,
    space_after: float = None,
    first_line_indent: float = None,
    left_indent: float = None,
    right_indent: float = None
) -> str:
    """
    在Word文档中创建自定义样式。
    
    Args:
        file_path: Word文档路径
        style_name: 样式名称
        style_type: 样式类型 (paragraph/character/table/list)
        based_on: 基于哪个样式（可选）
        font_name: 字体名称
        font_size: 字体大小（磅值）
        font_bold: 是否加粗
        font_italic: 是否斜体
        font_underline: 是否下划线
        font_color: 字体颜色（十六进制格式）
        alignment: 对齐方式
        line_spacing: 行间距
        space_before: 段前间距（磅值）
        space_after: 段后间距（磅值）
        first_line_indent: 首行缩进（厘米）
        left_indent: 左缩进（厘米）
        right_indent: 右缩进（厘米）
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法创建自定义样式，请先安装python-docx库"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 映射样式类型
    style_type_map = {
        "paragraph": WD_STYLE_TYPE.PARAGRAPH,
        "character": WD_STYLE_TYPE.CHARACTER,
        "table": WD_STYLE_TYPE.TABLE,
        "list": WD_STYLE_TYPE.LIST
    }
    
    if style_type not in style_type_map:
        return f"错误: 不支持的样式类型 '{style_type}'，可用选项: paragraph, character, table, list"
    
    # 映射对齐方式
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    try:
        doc = Document(file_path)
        
        # 检查样式名称是否已存在
        style_exists = False
        for style in doc.styles:
            if style.name == style_name:
                style_exists = True
                break
        
        # 创建新样式或获取现有样式
        if style_exists:
            return f"错误: 样式 '{style_name}' 已存在，请使用不同的名称或使用apply_style功能修改现有样式"
        else:
            style = doc.styles.add_style(style_name, style_type_map[style_type])
        
        # 设置基于哪个样式
        if based_on:
            try:
                style.base_style = doc.styles[based_on]
            except KeyError:
                return f"错误: 基础样式 '{based_on}' 不存在"
        
        # 设置字体属性
        if any([font_name, font_size, font_bold is not None, font_italic is not None, 
                font_underline is not None, font_color]):
                
            # 对于段落和字符样式，可以设置字体
            if style_type in ["paragraph", "character"]:
                font = style.font
                
                if font_name:
                    font.name = font_name
                    # 设置中文字体
                    try:
                        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    except:
                        pass
                
                if font_size:
                    font.size = Pt(font_size)
                
                if font_bold is not None:
                    font.bold = font_bold
                
                if font_italic is not None:
                    font.italic = font_italic
                
                if font_underline is not None:
                    font.underline = font_underline
                
                if font_color:
                    try:
                        # 解析十六进制颜色
                        color = font_color.lstrip('#')
                        rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                        font.color.rgb = RGBColor(*rgb)
                    except:
                        pass
        
        # 对于段落样式，可以设置段落格式
        if style_type == "paragraph":
            paragraph_format = style.paragraph_format
            
            if alignment and alignment in alignment_map:
                paragraph_format.alignment = alignment_map[alignment]
            
            if line_spacing is not None:
                paragraph_format.line_spacing = line_spacing
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            
            if space_before is not None:
                paragraph_format.space_before = Pt(space_before)
            
            if space_after is not None:
                paragraph_format.space_after = Pt(space_after)
            
            if first_line_indent is not None:
                paragraph_format.first_line_indent = Cm(first_line_indent)
            
            if left_indent is not None:
                paragraph_format.left_indent = Cm(left_indent)
            
            if right_indent is not None:
                paragraph_format.right_indent = Cm(right_indent)
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功在文档 {os.path.basename(file_path)} 中创建样式 '{style_name}'"
    
    except Exception as e:
        return f"创建自定义样式时出错: {str(e)}"

def apply_style(
    file_path: str,
    paragraph_indices: List[int],
    style_name: str,
    create_if_not_exists: bool = False,
    style_properties: dict = None
) -> str:
    """
    将指定样式应用到文档中的段落。
    
    Args:
        file_path: Word文档路径
        paragraph_indices: 段落索引列表
        style_name: 要应用的样式名称
        create_if_not_exists: 如果样式不存在是否创建
        style_properties: 创建新样式的属性（仅在create_if_not_exists为True时有效）
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法应用样式，请先安装python-docx库"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        doc = Document(file_path)
        
        # 检查样式是否存在
        style_exists = False
        for style in doc.styles:
            if style.name == style_name:
                style_exists = True
                break
        
        # 如果样式不存在且需要创建
        if not style_exists and create_if_not_exists:
            if style_properties is None:
                style_properties = {}
            
            # 创建默认段落样式
            style_properties["style_name"] = style_name
            style_properties["file_path"] = file_path
            
            # 调用创建样式函数
            create_result = create_custom_style(**style_properties)
            
            # 如果创建失败则返回错误
            if "错误" in create_result:
                return create_result
            
            # 重新加载文档以获取新创建的样式
            doc = Document(file_path)
        elif not style_exists:
            return f"错误: 样式 '{style_name}' 不存在，请先创建样式或设置create_if_not_exists=True"
        
        # 应用样式到指定段落
        success_count = 0
        invalid_indices = []
        
        for idx in paragraph_indices:
            # 检查段落索引是否有效
            if idx < 0 or idx >= len(doc.paragraphs):
                invalid_indices.append(idx)
                continue
            
            # 应用样式
            doc.paragraphs[idx].style = style_name
            success_count += 1
        
        # 保存文档
        doc.save(file_path)
        
        result_msg = f"成功将样式 '{style_name}' 应用到文档 {os.path.basename(file_path)} 中的 {success_count} 个段落"
        if invalid_indices:
            result_msg += f"，但有 {len(invalid_indices)} 个无效的段落索引: {invalid_indices}"
        
        return result_msg
    
    except Exception as e:
        return f"应用样式时出错: {str(e)}"

def export_document_styles(
    file_path: str,
    output_path: str = None,
    style_names: List[str] = None,  # 要导出的样式名称列表，为None表示导出所有
) -> str:
    """
    将文档中的样式导出到JSON文件。
    
    Args:
        file_path: Word文档路径
        output_path: 导出的JSON文件路径（不含扩展名），默认与文档同名
        style_names: 要导出的样式名称列表，None表示导出所有
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法导出样式，请先安装python-docx库"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 设置导出文件路径
    if output_path is None:
        # 使用与文档同名的路径，但修改扩展名为.json
        output_path = os.path.splitext(file_path)[0] + "_styles.json"
    elif not output_path.endswith('.json'):
        output_path += '_styles.json'
    
    try:
        doc = Document(file_path)
        
        # 收集样式信息
        style_info = []
        
        for style in doc.styles:
            # 如果指定了样式名称列表，只导出列表中的样式
            if style_names is not None and style.name not in style_names:
                continue
            
            # 创建样式信息字典
            style_data = {
                "name": style.name,
                "type": style.type,
                "properties": {}
            }
            
            # 记录基础样式
            if style.base_style:
                style_data["based_on"] = style.base_style.name
            
            # 添加字体属性
            if hasattr(style, 'font'):
                font_properties = {}
                
                if style.font.name:
                    font_properties["name"] = style.font.name
                
                if style.font.size:
                    font_properties["size"] = style.font.size.pt
                
                font_properties["bold"] = style.font.bold
                font_properties["italic"] = style.font.italic
                font_properties["underline"] = style.font.underline
                
                if style.font.color.rgb:
                    r, g, b = style.font.color.rgb.red, style.font.color.rgb.green, style.font.color.rgb.blue
                    font_properties["color"] = f"#{r:02x}{g:02x}{b:02x}"
                
                style_data["properties"]["font"] = font_properties
            
            # 添加段落格式属性
            if hasattr(style, 'paragraph_format'):
                para_properties = {}
                
                if hasattr(style.paragraph_format, 'alignment') and style.paragraph_format.alignment:
                    alignment_map_reverse = {
                        WD_ALIGN_PARAGRAPH.LEFT: "left",
                        WD_ALIGN_PARAGRAPH.RIGHT: "right",
                        WD_ALIGN_PARAGRAPH.CENTER: "center",
                        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
                    }
                    para_properties["alignment"] = alignment_map_reverse.get(style.paragraph_format.alignment, None)
                
                if style.paragraph_format.line_spacing:
                    para_properties["line_spacing"] = style.paragraph_format.line_spacing
                
                if style.paragraph_format.space_before:
                    para_properties["space_before"] = style.paragraph_format.space_before.pt
                
                if style.paragraph_format.space_after:
                    para_properties["space_after"] = style.paragraph_format.space_after.pt
                
                if style.paragraph_format.first_line_indent:
                    para_properties["first_line_indent"] = style.paragraph_format.first_line_indent.cm
                
                if style.paragraph_format.left_indent:
                    para_properties["left_indent"] = style.paragraph_format.left_indent.cm
                
                if style.paragraph_format.right_indent:
                    para_properties["right_indent"] = style.paragraph_format.right_indent.cm
                
                style_data["properties"]["paragraph_format"] = para_properties
            
            style_info.append(style_data)
        
        # 将样式信息保存到JSON文件
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(style_info, f, indent=4, ensure_ascii=False)
        
        count = len(style_info)
        return f"成功导出 {count} 个样式到 {output_path}"
    
    except Exception as e:
        return f"导出样式时出错: {str(e)}"

def import_document_styles(
    file_path: str,
    style_file_path: str,
    style_names: List[str] = None,  # 要导入的样式名称列表，为None表示导入所有
    overwrite_existing: bool = False
) -> str:
    """
    从JSON文件导入样式到Word文档。
    
    Args:
        file_path: Word文档路径
        style_file_path: 样式JSON文件路径
        style_names: 要导入的样式名称列表，None表示导入所有
        overwrite_existing: 是否覆盖现有样式
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法导入样式，请先安装python-docx库"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 确保样式文件存在
    if not os.path.exists(style_file_path):
        return f"错误: 样式文件 {style_file_path} 不存在"
    
    try:
        # 加载样式信息
        with open(style_file_path, 'r', encoding='utf-8') as f:
            style_info = json.load(f)
        
        doc = Document(file_path)
        
        # 获取文档中现有的样式名称
        existing_styles = [style.name for style in doc.styles]
        
        # 记录操作结果
        imported_count = 0
        skipped_count = 0
        failed_styles = []
        
        # 导入样式
        for style_data in style_info:
            style_name = style_data["name"]
            
            # 如果指定了样式名称列表，只导入列表中的样式
            if style_names is not None and style_name not in style_names:
                continue
            
            # 检查样式是否已存在
            if style_name in existing_styles and not overwrite_existing:
                skipped_count += 1
                continue
            
            # 准备样式参数
            style_params = {
                "file_path": file_path,
                "style_name": style_name
            }
            
            # 设置样式类型
            style_type_map = {
                1: "paragraph",
                2: "character",
                3: "table",
                4: "list"
            }
            style_params["style_type"] = style_type_map.get(style_data["type"], "paragraph")
            
            # 设置基础样式
            if "based_on" in style_data:
                style_params["based_on"] = style_data["based_on"]
            
            # 设置字体属性
            if "properties" in style_data and "font" in style_data["properties"]:
                font = style_data["properties"]["font"]
                
                if "name" in font:
                    style_params["font_name"] = font["name"]
                
                if "size" in font:
                    style_params["font_size"] = font["size"]
                
                if "bold" in font:
                    style_params["font_bold"] = font["bold"]
                
                if "italic" in font:
                    style_params["font_italic"] = font["italic"]
                
                if "underline" in font:
                    style_params["font_underline"] = font["underline"]
                
                if "color" in font:
                    style_params["font_color"] = font["color"]
            
            # 设置段落格式属性
            if "properties" in style_data and "paragraph_format" in style_data["properties"]:
                para_format = style_data["properties"]["paragraph_format"]
                
                if "alignment" in para_format:
                    style_params["alignment"] = para_format["alignment"]
                
                if "line_spacing" in para_format:
                    style_params["line_spacing"] = para_format["line_spacing"]
                
                if "space_before" in para_format:
                    style_params["space_before"] = para_format["space_before"]
                
                if "space_after" in para_format:
                    style_params["space_after"] = para_format["space_after"]
                
                if "first_line_indent" in para_format:
                    style_params["first_line_indent"] = para_format["first_line_indent"]
                
                if "left_indent" in para_format:
                    style_params["left_indent"] = para_format["left_indent"]
                
                if "right_indent" in para_format:
                    style_params["right_indent"] = para_format["right_indent"]
            
            try:
                # 创建新样式或更新现有样式
                if style_name in existing_styles and overwrite_existing:
                    # 删除现有样式
                    # 注意：python-docx不直接支持删除样式，这里我们通过创建新样式来覆盖
                    pass
                
                # 创建新样式
                result = create_custom_style(**style_params)
                
                if "成功" in result:
                    imported_count += 1
                else:
                    failed_styles.append(style_name)
            except Exception as e:
                failed_styles.append(style_name)
        
        result_msg = f"导入样式结果: 成功 {imported_count} 个, 跳过 {skipped_count} 个"
        if failed_styles:
            result_msg += f", 失败 {len(failed_styles)} 个: {failed_styles}"
        
        return result_msg
    
    except Exception as e:
        return f"导入样式时出错: {str(e)}"

def copy_style_between_documents(
    source_file_path: str,
    target_file_path: str,
    style_names: List[str],
    overwrite_existing: bool = False
) -> str:
    """
    在文档之间复制样式。
    
    Args:
        source_file_path: 源文档路径
        target_file_path: 目标文档路径
        style_names: 要复制的样式名称列表
        overwrite_existing: 是否覆盖目标文档中的现有样式
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法复制样式，请先安装python-docx库"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(source_file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        source_file_path = os.path.join(base_path, source_file_path)
    
    if not os.path.isabs(target_file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        target_file_path = os.path.join(base_path, target_file_path)
    
    # 确保源文件存在
    if not os.path.exists(source_file_path):
        return f"错误: 源文件 {source_file_path} 不存在"
    
    # 确保目标文件存在
    if not os.path.exists(target_file_path):
        return f"错误: 目标文件 {target_file_path} 不存在"
    
    try:
        # 创建临时样式文件
        temp_style_file = os.path.join(os.path.dirname(source_file_path), "_temp_styles_.json")
        
        # 从源文档导出指定样式
        export_result = export_document_styles(source_file_path, temp_style_file, style_names)
        
        # 检查导出是否成功
        if "错误" in export_result:
            # 清理临时文件
            if os.path.exists(temp_style_file):
                os.remove(temp_style_file)
            return f"从源文档导出样式时出错: {export_result}"
        
        # 将样式导入到目标文档
        import_result = import_document_styles(target_file_path, temp_style_file, style_names, overwrite_existing)
        
        # 清理临时文件
        if os.path.exists(temp_style_file):
            os.remove(temp_style_file)
        
        # 返回导入结果
        return f"将样式从 {os.path.basename(source_file_path)} 复制到 {os.path.basename(target_file_path)}: {import_result}"
    
    except Exception as e:
        # 清理临时文件
        if 'temp_style_file' in locals() and os.path.exists(temp_style_file):
            os.remove(temp_style_file)
        return f"复制样式时出错: {str(e)}" 