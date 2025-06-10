"""
Word Document Edit Operations

This module provides functions for editing Word document content including
paragraph editing, text replacement, and paragraph deletion.
"""
import os
import sys
from typing import Union, List


# 标记库是否已安装
docx_installed = True

# 尝试导入python-docx库，如果没有安装则标记为未安装但不退出
try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENTATION, WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("警告: 未检测到python-docx库，Word文档功能将不可用")
    print("请使用以下命令安装: pip install python-docx")
    docx_installed = False


def edit_paragraph_in_document(
    file_path: str,
    paragraph_index: int,
    new_text: str,
    save: bool = True,
    end_index: int = None,
    replacement_texts: list = None
) -> str:
    """
    编辑Word文档中指定段落或段落范围的文本内容。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        paragraph_index: 起始段落索引 (从0开始计数)
        new_text: 新的文本内容（用于单段落编辑）
        save: 是否保存更改，默认为True
        end_index: 结束段落索引（包含），如果为None则只编辑单个段落
        replacement_texts: 替换文本列表，用于批量编辑多个段落
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法编辑Word文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 确定编辑范围
        if end_index is None:
            # 单段落编辑模式
            start_idx = paragraph_index
            end_idx = paragraph_index
            edit_mode = "single"
        else:
            # 批量编辑模式
            start_idx = paragraph_index
            end_idx = end_index
            edit_mode = "batch"
        
        # 检查索引范围是否有效
        if start_idx < 0 or start_idx >= len(doc.paragraphs):
            return f"错误: 无效的起始段落索引 {start_idx}，文档共有 {len(doc.paragraphs)} 个段落"
        
        if end_idx < 0 or end_idx >= len(doc.paragraphs):
            return f"错误: 无效的结束段落索引 {end_idx}，文档共有 {len(doc.paragraphs)} 个段落"
        
        if start_idx > end_idx:
            return f"错误: 起始索引 {start_idx} 不能大于结束索引 {end_idx}"
        
        # 计算需要编辑的段落数量
        paragraph_count = end_idx - start_idx + 1
        modified_count = 0
        
        # 执行编辑操作
        if edit_mode == "single":
            # 单段落编辑
            paragraph = doc.paragraphs[start_idx]
            
            # 保存原始样式和格式
            original_style = paragraph.style
            original_alignment = paragraph.alignment
            
            # 清除现有内容
            for run in paragraph.runs:
                run.clear()
            
            # 确保段落内容被清除
            if paragraph.runs:
                paragraph.text = ""
            
            # 添加新内容
            run = paragraph.add_run(new_text)
            
            # 恢复原始样式和格式
            paragraph.style = original_style
            paragraph.alignment = original_alignment
            
            modified_count = 1
            
        else:
            # 批量编辑模式
            for i in range(start_idx, end_idx + 1):
                paragraph = doc.paragraphs[i]
                
                # 保存原始样式和格式
                original_style = paragraph.style
                original_alignment = paragraph.alignment
                
                # 清除现有内容
                for run in paragraph.runs:
                    run.clear()
                
                # 确保段落内容被清除
                if paragraph.runs:
                    paragraph.text = ""
                
                # 确定要使用的替换文本
                if replacement_texts and len(replacement_texts) > (i - start_idx):
                    # 使用对应的替换文本
                    replacement_text = replacement_texts[i - start_idx]
                elif replacement_texts and len(replacement_texts) == 1:
                    # 如果只提供了一个替换文本，所有段落都使用这个文本
                    replacement_text = replacement_texts[0]
                else:
                    # 使用默认的new_text
                    replacement_text = new_text
                
                # 添加新内容
                run = paragraph.add_run(replacement_text)
                
                # 恢复原始样式和格式
                paragraph.style = original_style
                paragraph.alignment = original_alignment
                
                modified_count += 1
        
        # 保存文档
        if save:
            doc.save(file_path)
        
        # 返回结果信息
        if edit_mode == "single":
            return f"成功编辑文档 {os.path.basename(file_path)} 第 {paragraph_index+1} 段落的内容"
        else:
            return f"成功编辑文档 {os.path.basename(file_path)} 第 {start_idx+1} 到第 {end_idx+1} 段落，共修改 {modified_count} 个段落"
            
    except Exception as e:
        return f"编辑Word文档内容时出错: {str(e)}"



def find_and_replace_text(
    file_path: str,
    find_text: str,
    replace_text: str,
    match_case: bool = False,
    match_whole_word: bool = False,
    save: bool = True
) -> str:
    """
    在Word文档中查找并替换文本。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        find_text: 要查找的文本
        replace_text: 替换为的文本
        match_case: 是否区分大小写，默认为False
        match_whole_word: 是否匹配整个单词，默认为False
        save: 是否保存更改，默认为True
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法在Word文档中查找替换文本，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 使用python-docx的方式（更可靠）
        doc = Document(file_path)
        replace_count = 0
        
        # 遍历所有段落和所有run
        for paragraph in doc.paragraphs:
            # 获取段落的完整文本
            full_text = paragraph.text
            if not match_case:
                search_text = find_text.lower()
                full_text_lower = full_text.lower()
            else:
                search_text = find_text
                full_text_lower = full_text
            
            # 如果段落中包含要查找的文本
            if search_text in full_text_lower:
                # 清除所有runs
                for run in paragraph.runs:
                    run.clear()
                
                # 如果需要区分大小写
                if match_case:
                    # 直接替换
                    new_text = full_text.replace(find_text, replace_text)
                    replace_count += full_text.count(find_text)
                else:
                    # 不区分大小写的替换
                    current_pos = 0
                    new_text = ""
                    while True:
                        # 在剩余文本中查找目标字符串
                        pos = full_text_lower[current_pos:].find(search_text)
                        if pos == -1:
                            # 没有找到更多匹配，添加剩余文本
                            new_text += full_text[current_pos:]
                            break
                        
                        # 添加匹配位置之前的文本
                        new_text += full_text[current_pos:current_pos + pos]
                        # 添加替换文本
                        new_text += replace_text
                        # 更新位置
                        current_pos += pos + len(find_text)
                        replace_count += 1
                
                # 添加新的run，包含替换后的文本
                paragraph.add_run(new_text)
        
        # 遍历所有表格单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 获取段落的完整文本
                        full_text = paragraph.text
                        if not match_case:
                            search_text = find_text.lower()
                            full_text_lower = full_text.lower()
                        else:
                            search_text = find_text
                            full_text_lower = full_text
                        
                        # 如果段落中包含要查找的文本
                        if search_text in full_text_lower:
                            # 清除所有runs
                            for run in paragraph.runs:
                                run.clear()
                            
                            # 如果需要区分大小写
                            if match_case:
                                # 直接替换
                                new_text = full_text.replace(find_text, replace_text)
                                replace_count += full_text.count(find_text)
                            else:
                                # 不区分大小写的替换
                                current_pos = 0
                                new_text = ""
                                while True:
                                    # 在剩余文本中查找目标字符串
                                    pos = full_text_lower[current_pos:].find(search_text)
                                    if pos == -1:
                                        # 没有找到更多匹配，添加剩余文本
                                        new_text += full_text[current_pos:]
                                        break
                                    
                                    # 添加匹配位置之前的文本
                                    new_text += full_text[current_pos:current_pos + pos]
                                    # 添加替换文本
                                    new_text += replace_text
                                    # 更新位置
                                    current_pos += pos + len(find_text)
                                    replace_count += 1
                            
                            # 添加新的run，包含替换后的文本
                            paragraph.add_run(new_text)
        
        # 保存文档
        if save:
            doc.save(file_path)
        
        return f"成功在文档 {os.path.basename(file_path)} 中替换了 {replace_count} 处文本"
    
    except Exception as e:
        return f"在Word文档中查找替换文本时出错: {str(e)}"


def delete_paragraph(
    file_path: str,
    paragraph_index: Union[int, List[int]],
    save: bool = True
) -> str:
    """
    删除Word文档中指定的段落或多个段落。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        paragraph_index: 要删除的段落索引 (从0开始计数) 或索引列表
        save: 是否保存更改，默认为True
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法编辑Word文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 处理输入参数，统一转换为列表格式
        if isinstance(paragraph_index, int):
            paragraph_indices = [paragraph_index]
            batch_mode = False
        else:
            paragraph_indices = paragraph_index
            batch_mode = True
        
        # 为了避免删除段落后索引变化导致的问题，先按索引降序排序
        sorted_indices = sorted(paragraph_indices, reverse=True)
        
        # 记录成功和失败的删除数量
        success_count = 0
        invalid_indices = []
        
        # 批量删除段落
        for idx in sorted_indices:
            # 检查段落索引是否有效
            if idx < 0 or idx >= len(doc.paragraphs):
                invalid_indices.append(idx)
                continue
            
            try:
                # 获取要删除的段落
                paragraph = doc.paragraphs[idx]
                
                # 删除段落
                p = paragraph._element
                p.getparent().remove(p)
                
                # 删除对象的引用
                paragraph._p = None
                paragraph._element = None
                
                success_count += 1
                
            except Exception as e:
                invalid_indices.append(idx)
        
        # 保存文档
        if save:
            doc.save(file_path)
        
        # 生成结果消息
        if batch_mode:
            result_msg = f"成功从文档 {os.path.basename(file_path)} 中删除 {success_count} 个段落"
            if invalid_indices:
                result_msg += f"，但有 {len(invalid_indices)} 个段落删除失败，索引: {invalid_indices}"
        else:
            if success_count > 0:
                result_msg = f"成功从文档 {os.path.basename(file_path)} 中删除第 {paragraph_indices[0]+1} 段落"
            else:
                result_msg = f"删除失败: 无效的段落索引 {paragraph_indices[0]}，文档共有 {len(doc.paragraphs)} 个段落"
        
        return result_msg
        
    except Exception as e:
        return f"删除Word文档段落时出错: {str(e)}"

