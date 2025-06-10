"""
批量段落操作模块 - 集成文本添加、格式设置和间距设置功能
支持批量操作以提高效率
"""

import os
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def batch_add_formatted_paragraphs(
    file_path: str,
    paragraphs_data: List[Dict[str, Any]],
    output_path: Optional[str] = None
) -> str:
    """
    批量添加格式化段落到Word文档
    
    Args:
        file_path: Word文档路径
        paragraphs_data: 段落数据列表，每个元素包含：
            - text: 文本内容
            - is_heading: 是否为标题（可选）
            - heading_level: 标题级别（可选）
            - alignment: 对齐方式（可选）
            - insert_position: 插入位置（可选）
            - font_name: 字体名称（可选）
            - font_size: 字体大小（可选）
            - bold: 是否加粗（可选）
            - italic: 是否斜体（可选）
            - underline: 是否下划线（可选）
            - font_color: 字体颜色（可选）
            - before_spacing: 段前间距（可选）
            - after_spacing: 段后间距（可选）
            - line_spacing: 行间距（可选）
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        doc = Document(file_path)
        success_count = 0
        failed_operations = []
        
        # 批量处理段落
        for i, para_data in enumerate(paragraphs_data):
            try:
                # 获取基本参数
                text = para_data.get('text', '')
                is_heading = para_data.get('is_heading', False)
                heading_level = para_data.get('heading_level', 1)
                alignment = para_data.get('alignment', 'left')
                insert_position = para_data.get('insert_position', -1)
                
                # 创建段落或标题
                if is_heading:
                    paragraph = doc.add_heading(text, level=heading_level)
                else:
                    paragraph = doc.add_paragraph(text)
                
                # 应用对齐方式
                alignment_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if alignment.lower() in alignment_map:
                    paragraph.alignment = alignment_map[alignment.lower()]
                
                # 应用文本格式
                _apply_text_formatting(paragraph, para_data)
                
                # 应用段落间距
                _apply_paragraph_spacing(paragraph, para_data)
                
                # 处理插入位置
                if insert_position != -1 and insert_position < len(doc.paragraphs) - 1:
                    target_paragraph = doc.paragraphs[insert_position]
                    new_p = paragraph._p
                    new_p.getparent().remove(new_p)
                    target_paragraph._p.addnext(new_p)
                
                success_count += 1
                
            except Exception as e:
                failed_operations.append((i, str(e)))
        
        # 保存文档
        doc.save(file_path)
        
        result_msg = f"成功批量添加 {success_count} 个格式化段落到文档 {os.path.basename(file_path)}"
        if failed_operations:
            result_msg += f"，但有 {len(failed_operations)} 个操作失败"
        
        return result_msg
        
    except Exception as e:
        return f"批量添加格式化段落时出错: {str(e)}"

def batch_format_paragraphs(
    file_path: str,
    format_operations: List[Dict[str, Any]],
    output_path: Optional[str] = None
) -> str:
    """
    批量格式化指定段落
    
    Args:
        file_path: Word文档路径
        format_operations: 格式化操作列表，每个元素包含：
            - paragraph_indices: 段落索引列表
            - font_name: 字体名称（可选）
            - font_size: 字体大小（可选）
            - bold: 是否加粗（可选）
            - italic: 是否斜体（可选）
            - underline: 是否下划线（可选）
            - font_color: 字体颜色（可选）
            - highlight_color: 高亮颜色（可选）
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        doc = Document(file_path)
        total_formatted = 0
        failed_operations = []
        
        # 批量处理格式化操作
        for i, operation in enumerate(format_operations):
            try:
                paragraph_indices = operation.get('paragraph_indices', [])
                
                for paragraph_index in paragraph_indices:
                    if 0 <= paragraph_index < len(doc.paragraphs):
                        paragraph = doc.paragraphs[paragraph_index]
                        _apply_text_formatting(paragraph, operation)
                        total_formatted += 1
                    else:
                        failed_operations.append((i, f"无效的段落索引: {paragraph_index}"))
                        
            except Exception as e:
                failed_operations.append((i, str(e)))
        
        # 保存文档
        doc.save(file_path)
        
        result_msg = f"成功批量格式化 {total_formatted} 个段落"
        if failed_operations:
            result_msg += f"，但有 {len(failed_operations)} 个操作失败"
        
        return result_msg
        
    except Exception as e:
        return f"批量格式化段落时出错: {str(e)}"

def batch_set_paragraph_spacing(
    file_path: str,
    spacing_operations: List[Dict[str, Any]],
    output_path: Optional[str] = None
) -> str:
    """
    批量设置段落间距
    
    Args:
        file_path: Word文档路径
        spacing_operations: 间距设置操作列表，每个元素包含：
            - paragraph_indices: 段落索引列表
            - before_spacing: 段前间距（可选）
            - after_spacing: 段后间距（可选）
            - line_spacing: 行间距（可选）
            - line_spacing_rule: 行间距规则（可选）
        output_path: 输出路径，如果为None则从环境变量获取
    
    Returns:
        操作结果信息
    """
    # 处理文件路径
    if output_path is None:
        output_path = os.environ.get('OFFICE_EDIT_PATH')
        if not output_path:
            output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(output_path, file_path)
    
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        doc = Document(file_path)
        total_processed = 0
        failed_operations = []
        
        # 批量处理间距设置操作
        for i, operation in enumerate(spacing_operations):
            try:
                paragraph_indices = operation.get('paragraph_indices', [])
                
                for paragraph_index in paragraph_indices:
                    if 0 <= paragraph_index < len(doc.paragraphs):
                        paragraph = doc.paragraphs[paragraph_index]
                        _apply_paragraph_spacing(paragraph, operation)
                        total_processed += 1
                    else:
                        failed_operations.append((i, f"无效的段落索引: {paragraph_index}"))
                        
            except Exception as e:
                failed_operations.append((i, str(e)))
        
        # 保存文档
        doc.save(file_path)
        
        result_msg = f"成功批量设置 {total_processed} 个段落的间距"
        if failed_operations:
            result_msg += f"，但有 {len(failed_operations)} 个操作失败"
        
        return result_msg
        
    except Exception as e:
        return f"批量设置段落间距时出错: {str(e)}"

def _apply_text_formatting(paragraph, format_data: Dict[str, Any]):
    """应用文本格式"""
    font_name = format_data.get('font_name')
    font_size = format_data.get('font_size')
    bold = format_data.get('bold', False)
    italic = format_data.get('italic', False)
    underline = format_data.get('underline', False)
    font_color = format_data.get('font_color')
    highlight_color = format_data.get('highlight_color')
    
    # 确保段落有run
    if len(paragraph.runs) == 0:
        original_text = paragraph.text
        for child in list(paragraph._element):
            paragraph._element.remove(child)
        paragraph.add_run(original_text)
    
    # 应用格式到所有runs
    for run in paragraph.runs:
        if font_name:
            run.font.name = font_name
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
        
        if font_size:
            run.font.size = Pt(font_size)
        
        run.font.bold = bold
        run.font.italic = italic
        run.font.underline = underline
        
        # 设置字体颜色
        if font_color:
            try:
                if font_color.startswith("#"):
                    font_color = font_color[1:]
                r = int(font_color[0:2], 16)
                g = int(font_color[2:4], 16)
                b = int(font_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                pass
        
        # 设置高亮颜色
        if highlight_color:
            highlight_color_map = {
                "yellow": "FFFF00", "green": "00FF00", "blue": "0000FF",
                "red": "FF0000", "pink": "FFC0CB", "turquoise": "40E0D0"
            }
            if highlight_color.lower() in highlight_color_map:
                shading_elm = OxmlElement('w:shd')
                color_value = highlight_color_map[highlight_color.lower()]
                shading_elm.set(qn('w:fill'), color_value)
                run._element.get_or_add_rPr().append(shading_elm)

def _apply_paragraph_spacing(paragraph, spacing_data: Dict[str, Any]):
    """应用段落间距"""
    before_spacing = spacing_data.get('before_spacing')
    after_spacing = spacing_data.get('after_spacing')
    line_spacing = spacing_data.get('line_spacing')
    line_spacing_rule = spacing_data.get('line_spacing_rule', 'multiple')
    
    # 设置段前间距
    if before_spacing is not None:
        paragraph.paragraph_format.space_before = Pt(before_spacing)
    
    # 设置段后间距
    if after_spacing is not None:
        paragraph.paragraph_format.space_after = Pt(after_spacing)
    
    # 设置行间距
    if line_spacing is not None:
        spacing_rule_map = {
            "multiple": WD_LINE_SPACING.MULTIPLE,
            "exact": WD_LINE_SPACING.EXACTLY,
            "atLeast": WD_LINE_SPACING.AT_LEAST
        }
        
        if line_spacing_rule in spacing_rule_map:
            paragraph.paragraph_format.line_spacing_rule = spacing_rule_map[line_spacing_rule]
            
            if line_spacing_rule == "multiple":
                paragraph.paragraph_format.line_spacing = line_spacing
            else:
                paragraph.paragraph_format.line_spacing = Pt(line_spacing)
