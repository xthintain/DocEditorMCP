"""
高级格式效果模块 - 实现文本框、艺术字、首字下沉等高级格式功能
"""

import os
from typing import List, Dict, Any, Optional, Union, Tuple
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

# 检查docx库安装状态
try:
    import docx
    docx_installed = True
except ImportError:
    docx_installed = False

# 尝试导入pywin32，用于某些高级功能
try:
    import win32com.client
    import pythoncom
    win32com_installed = True
except ImportError:
    win32com_installed = False

def add_text_box(
    file_path: str,
    text: str,
    width: float = 10.0,  # 厘米
    height: float = 5.0,  # 厘米
    position: str = "center",  # center, left, right
    border_style: str = "single",  # single, double, none
    border_color: str = None,  # 十六进制颜色，如 "#000000"
    fill_color: str = None,  # 十六进制颜色，如 "#FFFFFF"
    font_name: str = None,
    font_size: float = None,
    font_bold: bool = False,
    font_italic: bool = False,
    font_color: str = None,  # 十六进制颜色，如 "#000000"
    paragraph_index: int = -1  # 插入位置，-1表示文档末尾
) -> str:
    """
    在Word文档中添加文本框。
    
    Args:
        file_path: Word文档路径
        text: 文本框内容
        width: 文本框宽度（厘米）
        height: 文本框高度（厘米）
        position: 位置 (center/left/right)
        border_style: 边框样式 (single/double/none)
        border_color: 边框颜色 (十六进制)
        fill_color: 填充颜色 (十六进制)
        font_name: 字体名称
        font_size: 字体大小（磅值）
        font_bold: 是否加粗
        font_italic: 是否斜体
        font_color: 字体颜色 (十六进制)
        paragraph_index: 插入位置段落索引
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法创建文本框，请先安装python-docx库"
    
    # 检查是否支持高级文本框功能
    if not win32com_installed:
        return "错误: 文本框功能需要pywin32支持，请先安装: pip install pywin32"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 使用COM接口添加文本框（这是最可靠的方式）
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(file_path)
        
        # 确定插入位置
        if paragraph_index == -1 or paragraph_index >= doc.Paragraphs.Count:
            # 移动到文档末尾
            doc.ActiveWindow.Selection.EndKey(Unit=6)  # 6 = wdStory
        else:
            # 移动到指定段落
            if paragraph_index < 0:
                paragraph_index = 0
            
            para = doc.Paragraphs[paragraph_index + 1]  # COM对象从1开始计数
            para.Range.Select()
        
        # 插入文本框
        shape = doc.Shapes.AddTextbox(
            Orientation=0,  # 0 = msoTextOrientationHorizontal
            Left=0,
            Top=0,
            Width=Cm(width).inches * 72,  # 转换为点数
            Height=Cm(height).inches * 72  # 转换为点数
        )
        
        # 设置位置
        if position == "center":
            shape.RelativeHorizontalPosition = 0  # 0 = wdRelativeHorizontalPositionMargin
            shape.RelativeVerticalPosition = 0    # 0 = wdRelativeVerticalPositionMargin
            shape.Left = 2  # 2 = wdShapeCenter
            shape.Top = 2   # 2 = wdShapeCenter
        elif position == "left":
            shape.RelativeHorizontalPosition = 0  # 0 = wdRelativeHorizontalPositionMargin
            shape.Left = 0  # 0 = wdShapeLeft
        elif position == "right":
            shape.RelativeHorizontalPosition = 0  # 0 = wdRelativeHorizontalPositionMargin
            shape.Left = 1  # 1 = wdShapeRight
        
        # 设置边框样式
        if border_style == "single":
            shape.Line.Weight = 0.75
            shape.Line.Style = 1  # 1 = msoLineSolid
        elif border_style == "double":
            shape.Line.Weight = 1.5
            shape.Line.Style = 1  # 1 = msoLineSolid
        elif border_style == "none":
            shape.Line.Visible = False
        
        # 设置边框颜色
        if border_color:
            color = border_color.lstrip('#')
            rgb_int = int(color, 16)
            shape.Line.ForeColor.RGB = rgb_int
        
        # 设置填充颜色
        if fill_color:
            color = fill_color.lstrip('#')
            rgb_int = int(color, 16)
            shape.Fill.ForeColor.RGB = rgb_int
        else:
            # 默认透明填充
            shape.Fill.Visible = False
        
        # 添加文本到文本框
        shape.TextFrame.TextRange.Text = text
        
        # 设置字体属性
        text_range = shape.TextFrame.TextRange
        
        if font_name:
            text_range.Font.Name = font_name
        
        if font_size:
            text_range.Font.Size = font_size
        
        if font_bold:
            text_range.Font.Bold = True
        
        if font_italic:
            text_range.Font.Italic = True
        
        if font_color:
            color = font_color.lstrip('#')
            rgb_int = int(color, 16)
            text_range.Font.Color.RGB = rgb_int
        
        # 保存并关闭文档
        doc.Save()
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        
        return f"成功添加文本框到文档 {os.path.basename(file_path)}"
    
    except Exception as e:
        if 'word' in locals():
            try:
                if 'doc' in locals():
                    doc.Close(SaveChanges=False)
                word.Quit()
            except:
                pass
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return f"添加文本框时出错: {str(e)}"

def add_drop_cap(
    file_path: str,
    paragraph_index: int,
    dropped_lines: int = 2,  # 下沉的行数
    font_name: str = None,
    font_color: str = None,  # 十六进制颜色，如 "#000000"
) -> str:
    """
    为Word文档中的段落添加首字下沉效果。
    
    Args:
        file_path: Word文档路径
        paragraph_index: 要添加首字下沉效果的段落索引
        dropped_lines: 下沉的行数（1-10）
        font_name: 字体名称
        font_color: 字体颜色（十六进制）
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法添加首字下沉，请先安装python-docx库"
    
    # 检查是否支持高级文本格式功能
    if not win32com_installed:
        return "错误: 首字下沉功能需要pywin32支持，请先安装: pip install pywin32"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 验证参数
    if dropped_lines < 1 or dropped_lines > 10:
        return "错误: 下沉行数必须在1到10之间"
    
    try:
        # 使用COM接口添加首字下沉（这是最可靠的方式）
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(file_path)
        
        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= doc.Paragraphs.Count:
            doc.Close(SaveChanges=False)
            word.Quit()
            pythoncom.CoUninitialize()
            return f"错误: 段落索引 {paragraph_index} 超出范围，文档共有 {doc.Paragraphs.Count} 个段落"
        
        # 获取段落并应用首字下沉
        para = doc.Paragraphs[paragraph_index + 1]  # COM对象从1开始计数
        
        # 确保段落有内容
        if len(para.Range.Text.strip()) == 0:
            doc.Close(SaveChanges=False)
            word.Quit()
            pythoncom.CoUninitialize()
            return f"错误: 段落 {paragraph_index} 没有内容，无法添加首字下沉"
        
        # 选择段落的第一个字符
        first_char = para.Range.Characters[1]
        first_char.Select()
        
        # 应用首字下沉
        word.Selection.Font.DropCap.Position = 1  # 1 = wdDropNormal
        word.Selection.Font.DropCap.LinesToDrop = dropped_lines
        
        # 设置字体属性
        if font_name:
            word.Selection.Font.Name = font_name
        
        if font_color:
            color = font_color.lstrip('#')
            rgb_int = int(color, 16)
            word.Selection.Font.Color = rgb_int
        
        # 保存并关闭文档
        doc.Save()
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        
        return f"成功为文档 {os.path.basename(file_path)} 的第 {paragraph_index} 段添加首字下沉效果"
    
    except Exception as e:
        if 'word' in locals():
            try:
                if 'doc' in locals():
                    doc.Close(SaveChanges=False)
                word.Quit()
            except:
                pass
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return f"添加首字下沉效果时出错: {str(e)}"

def add_word_art(
    file_path: str,
    text: str,
    style: int = 1,  # 1-47 对应Word中的艺术字样式
    size: float = 36.0,  # 磅值
    fill_color: str = None,  # 十六进制颜色，如 "#000000"
    outline_color: str = None,  # 十六进制颜色，如 "#FFFFFF"
    paragraph_index: int = -1  # 插入位置，-1表示文档末尾
) -> str:
    """
    在Word文档中添加艺术字。
    
    Args:
        file_path: Word文档路径
        text: 艺术字文本内容
        style: 艺术字样式编号（1-47）
        size: 艺术字大小（磅值）
        fill_color: 填充颜色（十六进制）
        outline_color: 轮廓颜色（十六进制）
        paragraph_index: 插入位置段落索引
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法添加艺术字，请先安装python-docx库"
    
    # 检查是否支持高级文本格式功能
    if not win32com_installed:
        return "错误: 艺术字功能需要pywin32支持，请先安装: pip install pywin32"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 验证参数
    if style < 1 or style > 47:
        return "错误: 艺术字样式必须在1到47之间"
    
    try:
        # 使用COM接口添加艺术字
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(file_path)
        
        # 确定插入位置
        if paragraph_index == -1 or paragraph_index >= doc.Paragraphs.Count:
            # 移动到文档末尾
            doc.ActiveWindow.Selection.EndKey(Unit=6)  # 6 = wdStory
        else:
            # 移动到指定段落
            if paragraph_index < 0:
                paragraph_index = 0
            
            para = doc.Paragraphs[paragraph_index + 1]  # COM对象从1开始计数
            para.Range.Select()
        
        # 插入艺术字
        art = word.Selection.InlineShapes.AddWordArt(
            PresetWordArtEffect=style,
            Text=text,
            FontName="Arial",
            FontSize=size,
            Bold=False,
            Italic=False
        )
        
        # 设置填充颜色
        if fill_color:
            color = fill_color.lstrip('#')
            rgb_int = int(color, 16)
            art.TextEffect.Fill.ForeColor.RGB = rgb_int
        
        # 设置轮廓颜色
        if outline_color:
            color = outline_color.lstrip('#')
            rgb_int = int(color, 16)
            art.TextEffect.Line.ForeColor.RGB = rgb_int
        
        # 保存并关闭文档
        doc.Save()
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        
        return f"成功添加艺术字到文档 {os.path.basename(file_path)}"
    
    except Exception as e:
        if 'word' in locals():
            try:
                if 'doc' in locals():
                    doc.Close(SaveChanges=False)
                word.Quit()
            except:
                pass
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return f"添加艺术字时出错: {str(e)}"

def add_custom_bullets(
    file_path: str,
    paragraph_indices: List[int],
    bullet_style: str = "disc",  # disc, circle, square, number, custom
    custom_symbol: str = None,
    font_name: str = None,
    font_color: str = None,  # 十六进制颜色，如 "#000000"
) -> str:
    """
    为Word文档中的段落添加自定义项目符号。
    
    Args:
        file_path: Word文档路径
        paragraph_indices: 要添加项目符号的段落索引列表
        bullet_style: 项目符号样式 (disc/circle/square/number/custom)
        custom_symbol: 自定义符号（仅当bullet_style为custom时有效）
        font_name: 符号字体名称
        font_color: 符号颜色（十六进制）
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法添加自定义项目符号，请先安装python-docx库"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 映射符号样式
    symbol_map = {
        "disc": "•",
        "circle": "○", 
        "square": "■",
        "number": "1."
    }
    
    # 如果是custom样式，必须提供自定义符号
    if bullet_style == "custom" and not custom_symbol:
        return "错误: 使用自定义项目符号样式时必须提供custom_symbol"
    
    try:
        # 尝试使用Word COM对象（功能最完整）
        if win32com_installed:
            try:
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(file_path)
                
                success_count = 0
                invalid_indices = []
                
                for idx in paragraph_indices:
                    # 检查段落索引是否有效
                    if idx < 0 or idx >= doc.Paragraphs.Count:
                        invalid_indices.append(idx)
                        continue
                    
                    # 获取段落并应用项目符号
                    para = doc.Paragraphs[idx + 1]  # COM对象从1开始计数
                    para.Range.Select()
                    
                    if bullet_style == "number":
                        # 使用编号列表
                        word.Selection.Range.ListFormat.ApplyNumberDefault()
                    elif bullet_style == "custom" or bullet_style in symbol_map:
                        # 使用项目符号
                        word.Selection.Range.ListFormat.ApplyBulletDefault()
                        
                        # 如果是自定义样式，设置自定义符号
                        if bullet_style == "custom":
                            word.Selection.Range.ListFormat.ListTemplate.ListLevels(1).NumberFormat = custom_symbol + " "
                        else:
                            word.Selection.Range.ListFormat.ListTemplate.ListLevels(1).NumberFormat = symbol_map[bullet_style] + " "
                        
                        # 设置字体
                        if font_name:
                            word.Selection.Range.ListFormat.ListTemplate.ListLevels(1).Font.Name = font_name
                        
                        # 设置颜色
                        if font_color:
                            color = font_color.lstrip('#')
                            rgb_int = int(color, 16)
                            word.Selection.Range.ListFormat.ListTemplate.ListLevels(1).Font.Color = rgb_int
                    
                    success_count += 1
                
                # 保存并关闭文档
                doc.Save()
                doc.Close()
                word.Quit()
                pythoncom.CoUninitialize()
                
                result_msg = f"成功为文档 {os.path.basename(file_path)} 中的 {success_count} 个段落添加项目符号"
                if invalid_indices:
                    result_msg += f"，但有 {len(invalid_indices)} 个无效的段落索引: {invalid_indices}"
                
                return result_msg
            
            except Exception as e:
                if 'word' in locals():
                    try:
                        if 'doc' in locals():
                            doc.Close(SaveChanges=False)
                        word.Quit()
                    except:
                        pass
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass
                return f"添加项目符号时出错: {str(e)}"
        
        else:
            # 使用python-docx添加基本的项目符号（功能有限）
            doc = Document(file_path)
            
            success_count = 0
            invalid_indices = []
            
            for idx in paragraph_indices:
                # 检查段落索引是否有效
                if idx < 0 or idx >= len(doc.paragraphs):
                    invalid_indices.append(idx)
                    continue
                
                # 获取段落并应用项目符号
                paragraph = doc.paragraphs[idx]
                
                # 使用python-docx添加简单项目符号（功能有限）
                pPr = paragraph._p.get_or_add_pPr()
                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), "ListBullet")
                pPr.append(pStyle)
                
                # 尝试添加项目符号（注：python-docx对项目符号的支持有限）
                numPr = OxmlElement("w:numPr")
                numId = OxmlElement("w:numId")
                numId.set(qn("w:val"), "1")
                numPr.append(numId)
                pPr.append(numPr)
                
                success_count += 1
            
            # 保存文档
            doc.save(file_path)
            
            result_msg = f"成功为文档 {os.path.basename(file_path)} 中的 {success_count} 个段落添加基本项目符号"
            if invalid_indices:
                result_msg += f"，但有 {len(invalid_indices)} 个无效的段落索引: {invalid_indices}"
            
            result_msg += "\n注意：使用python-docx实现的项目符号功能有限，建议安装pywin32获得完整支持。"
            
            return result_msg
    
    except Exception as e:
        return f"添加项目符号时出错: {str(e)}" 