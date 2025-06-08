"""
MCP Server for Word Document Operations

This server provides tools to create, edit and manage Word documents.
It's implemented using the Model Context Protocol (MCP) Python SDK.
"""

import os
import sys
import io
from mcp.server.fastmcp import FastMCP
from typing import Optional, List, Dict, Any, Union, Tuple

# 标记库是否已安装
docx_installed = True

# 尝试导入python-docx库，如果没有安装则标记为未安装但不退出
try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENTATION, WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("警告: 未检测到python-docx库，Word文档功能将不可用")
    print("请使用以下命令安装: pip install python-docx")
    docx_installed = False

# 尝试导入Pillow库，用于图片处理
pillow_installed = True
try:
    from PIL import Image
except ImportError:
    print("警告: 未检测到Pillow库，图片处理功能将受限")
    print("请使用以下命令安装: pip install Pillow")
    pillow_installed = False

# 创建一个MCP服务器，保持名称与配置文件一致
mcp = FastMCP("wordEditor", enable_standard_requests=True)

# 添加回原始的TXT文件创建功能，确保基本功能可用
@mcp.tool()
def create_empty_txt(filename: str) -> str:
    """
    在指定路径上创建一个空白的TXT文件。
    
    Args:
        filename: 要创建的文件名 (不需要包含.txt扩展名)
    
    Returns:
        包含操作结果的消息
    """
    # 确保文件名有.txt扩展名
    if not filename.lower().endswith('.txt'):
        filename += '.txt'
    
    # 从环境变量获取输出路径，如果未设置则使用默认桌面路径
    output_path = os.environ.get('OFFICE_EDIT_PATH')
    if not output_path:
        output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    # 创建完整的文件路径
    file_path = os.path.join(output_path, filename)
    
    try:
        # 创建输出目录（如果不存在）
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 创建空白文件
        with open(file_path, 'w', encoding='utf-8') as f:
            pass
        return f"成功在 {output_path} 创建了空白文件: {filename}"
    except Exception as e:
        return f"创建文件时出错: {str(e)}"

@mcp.tool()
def create_word_document(filename: str) -> str:
    """
    创建一个新的Word文档。
    
    Args:
        filename: 要创建的文件名 (不需要包含.docx扩展名)
    
    Returns:
        包含操作结果的消息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法创建Word文档，请先安装python-docx库: pip install python-docx"
    
    # 确保文件名有.docx扩展名
    if not filename.lower().endswith('.docx'):
        filename += '.docx'
    
    # 从环境变量获取输出路径，如果未设置则使用默认桌面路径
    output_path = os.environ.get('OFFICE_EDIT_PATH')  # 保持环境变量名不变，以兼容现有配置
    if not output_path:
        output_path = os.path.join(os.path.expanduser('~'), '桌面')
    
    # 创建完整的文件路径
    file_path = os.path.join(output_path, filename)
    
    try:
        # 创建输出目录（如果不存在）
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 创建新的Word文档
        doc = Document()
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功在 {output_path} 创建了Word文档: {filename}"
    except Exception as e:
        return f"创建Word文档时出错: {str(e)}"

@mcp.tool()
def open_and_read_word_document(file_path: str) -> str:
    """
    打开并读取Word文档的完整内容。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
    
    Returns:
        文档的完整内容
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法读取Word文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 提取文档基本信息
        paragraphs = [p.text for p in doc.paragraphs]
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        # 构建文档信息头
        doc_info = (
            f"文件名: {os.path.basename(file_path)}\n"
            f"段落数: {len(paragraphs)}\n"
            f"标题数: {len(headings)}\n\n"
        )
        
        # 构建完整文档内容，保留段落结构，并在每段前添加段落编号
        full_content = ""
        for i, p_text in enumerate(paragraphs):
            # 添加段落编号 (i) 在每段前
            full_content += f"[{i}] {p_text}\n"
        
        # 返回文档信息和完整内容
        return doc_info + full_content
    except Exception as e:
        return f"读取Word文档时出错: {str(e)}"


@mcp.tool()
def format_text_in_document(
    file_path: str,
    paragraph_index: int,
    font_name: str = None,
    font_size: int = None,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    font_color: str = None,
    highlight_color: str = None
) -> str:
    """
    设置Word文档中指定段落的文本格式。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        paragraph_index: 段落索引 (从0开始计数)
        font_name: 字体名称 (可选)
        font_size: 字体大小 (点数，可选)
        bold: 是否加粗
        italic: 是否斜体
        underline: 是否下划线
        font_color: 字体颜色 (十六进制RGB格式，如"#FF0000"表示红色，可选)
        highlight_color: 突出显示颜色 (可选，有效值: "yellow", "green", "blue", "red", "pink", "turquoise", "violet", "darkblue", "teal", "darkred", "darkgreen")
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法格式化Word文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 定义高亮颜色映射 (用于XML着色)
    highlight_color_map = {
        "yellow": "FFFF00",      
        "green": "00FF00",       
        "blue": "0000FF",        
        "red": "FF0000",         
        "pink": "FFC0CB",        
        "turquoise": "40E0D0",   
        "violet": "EE82EE",      
        "darkblue": "00008B",    
        "teal": "008080",        
        "darkred": "8B0000",     
        "darkgreen": "006400"    
    }
    
    # 校验高亮颜色
    if highlight_color and highlight_color.lower() not in highlight_color_map:
        return f"错误: 不支持的高亮颜色 '{highlight_color}'，可选值为: {', '.join(highlight_color_map.keys())}"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"错误: 无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落"
        
        # 获取指定的段落
        paragraph = doc.paragraphs[paragraph_index]
        
        # 检查段落是否有内容
        if not paragraph.text.strip():
            return f"警告: 段落 {paragraph_index+1} 为空或只包含空白字符，无法设置格式"
            
        # 检查段落是否有run，如果没有，添加一个run
        if len(paragraph.runs) == 0:
            # 保存原始文本
            original_text = paragraph.text
            # 清空段落
            for child in list(paragraph._element):
                paragraph._element.remove(child)
            # 添加新run
            paragraph.add_run(original_text)
        
        # 应用格式设置
        for run in paragraph.runs:
            if font_name:
                # 设置西文字体名
                run.font.name = font_name
                # 设置中文字体名（这是关键部分）
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
            
            if font_size:
                run.font.size = Pt(font_size)
            
            run.font.bold = bold
            run.font.italic = italic
            run.font.underline = underline
            
            # 设置字体颜色
            if font_color:
                try:
                    # 解析十六进制颜色值
                    if font_color.startswith("#"):
                        font_color = font_color[1:]
                    r = int(font_color[0:2], 16)
                    g = int(font_color[2:4], 16)
                    b = int(font_color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except ValueError:
                    return f"错误: 无效的字体颜色格式 '{font_color}'，请使用十六进制RGB格式，如 '#FF0000'"
            
            # 设置高亮颜色（通过XML方式）
            if highlight_color:
                shading_elm = OxmlElement('w:shd')
                color_value = highlight_color_map[highlight_color.lower()]
                shading_elm.set(qn('w:fill'), color_value)
                run._element.get_or_add_rPr().append(shading_elm)
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功设置文档 {os.path.basename(file_path)} 第 {paragraph_index+1} 段落的格式"
    
    except Exception as e:
        return f"设置Word文档格式时出错: {str(e)}"

@mcp.tool()
def set_paragraph_spacing(
    file_path: str,
    paragraph_index: int,
    before_spacing: float = None,
    after_spacing: float = None,
    line_spacing: float = None,
    line_spacing_rule: str = "multiple"
) -> str:
    """
    设置Word文档中指定段落的间距。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        paragraph_index: 段落索引 (从0开始计数)
        before_spacing: 段前间距 (磅，可选)
        after_spacing: 段后间距 (磅，可选)
        line_spacing: 行间距值 (当line_spacing_rule为"multiple"时为倍数，为"exact"时为磅值)
        line_spacing_rule: 行间距规则，可选值: "multiple"(倍数), "exact"(固定值), "atLeast"(最小值)
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法设置段落间距，请先安装python-docx库: pip install python-docx"
    
    # 打印接收到的参数，用于调试
    print(f"接收到的参数: file_path={file_path}, paragraph_index={paragraph_index}, "
          f"before_spacing={before_spacing}, after_spacing={after_spacing}, "
          f"line_spacing={line_spacing}, line_spacing_rule={line_spacing_rule}")
    
    # 检查参数是否为None
    if paragraph_index is None:
        return "错误: 必须提供段落索引(paragraph_index)参数"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 映射行间距规则
    spacing_rule_map = {
        "multiple": WD_LINE_SPACING.MULTIPLE,
        "exact": WD_LINE_SPACING.EXACTLY,
        "atLeast": WD_LINE_SPACING.AT_LEAST
    }
    
    if line_spacing_rule not in spacing_rule_map:
        return f"错误: 无效的行间距规则 '{line_spacing_rule}'，可选值为: multiple, exact, atLeast"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 确保paragraph_index是整数
        try:
            paragraph_index = int(paragraph_index)
        except (ValueError, TypeError):
            return f"错误: 段落索引必须是整数，收到的是: {paragraph_index}"
        
        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"错误: 无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落"
        
        # 获取指定的段落
        paragraph = doc.paragraphs[paragraph_index]
        
        # 设置段前间距
        if before_spacing is not None:
            paragraph.paragraph_format.space_before = Pt(before_spacing)
        
        # 设置段后间距
        if after_spacing is not None:
            paragraph.paragraph_format.space_after = Pt(after_spacing)
        
        # 设置行间距
        if line_spacing is not None:
            # 设置行间距规则
            paragraph.paragraph_format.line_spacing_rule = spacing_rule_map[line_spacing_rule]
            
            # 根据规则设置行间距值
            if line_spacing_rule == "multiple":
                paragraph.paragraph_format.line_spacing = line_spacing
            else:
                paragraph.paragraph_format.line_spacing = Pt(line_spacing)
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功设置文档 {os.path.basename(file_path)} 第 {paragraph_index+1} 段落的间距"
    except Exception as e:
        return f"设置段落间距时出错: {str(e)}"

@mcp.tool()
def insert_image(
    file_path: str,
    image_path: str,
    width: float = None,
    height: float = None,
    after_paragraph: int = -1
) -> str:
    """
    在Word文档中插入图片。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        image_path: 图片文件的完整路径或相对于输出目录的路径，支持本地图片
        width: 图片宽度（厘米，可选，如果不指定则保持原始比例）
        height: 图片高度（厘米，可选，如果不指定则保持原始比例）
        after_paragraph: 在指定段落后插入图片，-1表示文档末尾
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法插入图片，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 处理图片路径，同样支持相对路径
    if not os.path.isabs(image_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        image_path = os.path.join(base_path, image_path)
    
    # 检查图片文件是否存在
    if not os.path.exists(image_path):
        return f"错误: 图片文件 {image_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 检查指定段落是否有效
        if after_paragraph >= len(doc.paragraphs):
            return f"错误: 无效的段落索引 {after_paragraph}，文档共有 {len(doc.paragraphs)} 个段落"
        
        # 在指定位置插入图片
        if after_paragraph == -1:
            # 在文档末尾插入图片
            paragraph = doc.add_paragraph()
        else:
            # 在指定段落后插入新段落，然后插入图片
            paragraph = doc.paragraphs[after_paragraph]
        
        # 设置图片尺寸
        if width and height:
            run = paragraph.add_run()
            run.add_picture(image_path, width=Cm(width), height=Cm(height))
        elif width:
            run = paragraph.add_run()
            run.add_picture(image_path, width=Cm(width))
        elif height:
            run = paragraph.add_run()
            run.add_picture(image_path, height=Cm(height))
        else:
            run = paragraph.add_run()
            run.add_picture(image_path)
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功在文档 {os.path.basename(file_path)} 中插入图片 {os.path.basename(image_path)}"
    except Exception as e:
        return f"插入图片时出错: {str(e)}"

@mcp.tool()
def insert_table(
    file_path: str,
    rows: int,
    cols: int,
    data: List[List[str]] = None,
    after_paragraph: int = -1,
    style: str = "Table Grid"
) -> str:
    """
    在Word文档中插入表格。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        rows: 表格行数
        cols: 表格列数
        data: 表格内容，二维数组，每个元素对应一个单元格的内容（可选）
        after_paragraph: 在指定段落后插入表格，-1表示文档末尾
        style: 表格样式，默认为"Table Grid"
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法插入表格，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 校验参数
    if rows <= 0 or cols <= 0:
        return "错误: 表格行数和列数必须大于0"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 检查指定段落是否有效
        if after_paragraph >= len(doc.paragraphs):
            return f"错误: 无效的段落索引 {after_paragraph}，文档共有 {len(doc.paragraphs)} 个段落"
        
        # 在指定位置插入表格
        if after_paragraph == -1:
            # 在文档末尾插入表格
            table = doc.add_table(rows=rows, cols=cols)
        else:
            # 获取指定段落的位置
            paragraph = doc.paragraphs[after_paragraph]
            # 在段落后插入表格
            table = doc.add_table(rows=rows, cols=cols)
            # 移动表格到指定段落后
            paragraph._p.addnext(table._tbl)
        
        # 设置表格样式
        table.style = style
        
        # 如果提供了数据，填充表格内容
        if data:
            for i, row_data in enumerate(data):
                if i < rows:  # 确保不超出表格行数
                    for j, cell_data in enumerate(row_data):
                        if j < cols:  # 确保不超出表格列数
                            table.cell(i, j).text = str(cell_data)
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功在文档 {os.path.basename(file_path)} 中插入 {rows}x{cols} 的表格"
    except Exception as e:
        return f"插入表格时出错: {str(e)}"

@mcp.tool()
def edit_table_cell(
    file_path: str,
    table_index: int,
    row: int,
    col: int,
    text: str
) -> str:
    """
    编辑Word文档中表格的单元格内容。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        table_index: 表格索引 (从0开始计数)
        row: 行索引 (从0开始计数)
        col: 列索引 (从0开始计数)
        text: 单元格内容
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法编辑表格，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 检查表格索引是否有效
        if table_index < 0 or table_index >= len(doc.tables):
            return f"错误: 无效的表格索引 {table_index}，文档共有 {len(doc.tables)} 个表格"
        
        # 获取指定的表格
        table = doc.tables[table_index]
        
        # 检查行索引是否有效
        if row < 0 or row >= len(table.rows):
            return f"错误: 无效的行索引 {row}，表格共有 {len(table.rows)} 行"
        
        # 检查列索引是否有效
        if col < 0 or col >= len(table.columns):
            return f"错误: 无效的列索引 {col}，表格共有 {len(table.columns)} 列"
        
        # 编辑单元格内容
        table.cell(row, col).text = text
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功编辑文档 {os.path.basename(file_path)} 中第 {table_index+1} 个表格的单元格 ({row+1},{col+1})"
    except Exception as e:
        return f"编辑表格单元格时出错: {str(e)}"

@mcp.tool()
def save_document_as_pdf(file_path: str) -> str:
    """
    将Word文档保存为PDF格式。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法导出PDF，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 构建PDF文件路径
        pdf_path = os.path.splitext(file_path)[0] + ".pdf"
        
        # 尝试使用Microsoft Word COM对象导出PDF
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17表示PDF格式
            doc.Close()
            word.Quit()
            
            return f"成功将文档导出为PDF: {os.path.basename(pdf_path)}"
        
        except ImportError:
            # 如果没有win32com库或不是Windows系统，返回错误信息
            return "错误: 导出PDF功能需要在Windows系统上安装pywin32库，请使用以下命令安装: pip install pywin32"
    
    except Exception as e:
        return f"导出PDF时出错: {str(e)}"

@mcp.tool()
def save_document_as(file_path: str, output_format: str = "docx", new_filename: str = None) -> str:
    """
    将Word文档保存为指定格式。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        output_format: 输出格式，可选值: "docx", "doc", "pdf", "txt", "html"
        new_filename: 新文件名(不含扩展名)，如果不提供则使用原文件名
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法保存文档，请先安装python-docx库: pip install python-docx"
    
    # 检查格式是否支持
    supported_formats = ["docx", "doc", "pdf", "txt", "html"]
    if output_format.lower() not in supported_formats:
        return f"错误: 不支持的输出格式 '{output_format}'，可选值为: {', '.join(supported_formats)}"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 构建新文件路径
        original_basename = os.path.splitext(os.path.basename(file_path))[0]
        output_dirname = os.path.dirname(file_path)
        
        # 如果提供了新文件名，则使用新文件名
        if new_filename:
            output_basename = new_filename
        else:
            output_basename = original_basename
        
        # 创建新文件的完整路径
        output_path = os.path.join(output_dirname, f"{output_basename}.{output_format}")
        
        # 根据输出格式选择不同的处理方式
        if output_format.lower() == "pdf":
            # 使用前面定义的导出PDF功能
            return save_document_as_pdf(file_path)
        
        elif output_format.lower() in ["docx", "doc"]:
            try:
                # 尝试使用Microsoft Word COM对象保存
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # 使用数字格式指定不同的Word格式
                format_map = {
                    "docx": 16,  # wdFormatDocumentDefault (*.docx)
                    "doc": 0     # wdFormatDocument97 (*.doc)
                }
                
                doc = word.Documents.Open(file_path)
                doc.SaveAs(output_path, FileFormat=format_map[output_format.lower()])
                doc.Close()
                word.Quit()
                
                return f"成功将文档保存为 {output_format} 格式: {os.path.basename(output_path)}"
            
            except ImportError:
                # 如果是docx格式，我们可以使用python-docx直接保存
                if output_format.lower() == "docx":
                    doc = Document(file_path)
                    doc.save(output_path)
                    return f"成功将文档保存为 DOCX 格式: {os.path.basename(output_path)}"
                else:
                    return "错误: 保存为DOC格式需要在Windows系统上安装pywin32库"
        
        elif output_format.lower() == "txt":
            # 将文档转换为纯文本
            doc = Document(file_path)
            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return f"成功将文档保存为文本格式: {os.path.basename(output_path)}"
        
        elif output_format.lower() == "html":
            try:
                # 尝试使用Microsoft Word COM对象保存为HTML
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(file_path)
                doc.SaveAs(output_path, FileFormat=8)  # 8表示HTML格式
                doc.Close()
                word.Quit()
                
                return f"成功将文档保存为HTML格式: {os.path.basename(output_path)}"
            
            except ImportError:
                return "错误: 保存为HTML格式需要在Windows系统上安装pywin32库"
    
    except Exception as e:
        return f"保存文档时出错: {str(e)}"

@mcp.tool()
def close_document(file_path: str, save_changes: bool = True) -> str:
    """
    关闭Word文档，可选是否保存更改。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        save_changes: 是否保存更改，默认为True
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法关闭文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 尝试使用Microsoft Word COM对象关闭文档
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            
            # 检查文档是否已经打开
            doc_found = False
            for doc in word.Documents:
                if os.path.abspath(doc.FullName) == os.path.abspath(file_path):
                    if save_changes:
                        doc.Save()
                    doc.Close(SaveChanges=save_changes)
                    doc_found = True
                    break
            
            if not doc_found:
                # 如果文档未打开，则python-docx中没有明确的关闭方法
                # 这里我们模拟一个"关闭"操作，即打开文档后立即关闭
                doc = Document(file_path)
                if save_changes:
                    doc.save(file_path)
            
            pythoncom.CoUninitialize()
            
            return f"成功关闭文档: {os.path.basename(file_path)}" + (" 并保存更改" if save_changes else "")
        
        except ImportError:
            # 如果没有win32com库，使用python-docx的方式
            doc = Document(file_path)
            if save_changes:
                doc.save(file_path)
            
            # python-docx没有明确的关闭方法，垃圾收集器会处理
            return f"成功关闭文档: {os.path.basename(file_path)}" + (" 并保存更改" if save_changes else "")
    
    except Exception as e:
        return f"关闭文档时出错: {str(e)}"

@mcp.tool()
def edit_paragraph_in_document(
    file_path: str,
    paragraph_index: int,
    new_text: str,
    save: bool = True
) -> str:
    """
    编辑Word文档中指定段落的文本内容。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        paragraph_index: 段落索引 (从0开始计数)
        new_text: 新的文本内容
        save: 是否保存更改，默认为True
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法编辑Word文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"错误: 无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落"
        
        # 获取并编辑指定的段落
        paragraph = doc.paragraphs[paragraph_index]
        
        # 保存原始样式和格式
        original_style = paragraph.style
        original_alignment = paragraph.alignment
        
        # 更简单安全的替换方法：清除所有runs并添加新文本
        for run in paragraph.runs:
            run.clear()
        
        # 清空所有runs后，确保段落内容被清除
        if paragraph.runs:
            # 如果仍有runs，直接重新设置text属性
            paragraph.text = ""
        
        # 添加新内容
        run = paragraph.add_run(new_text)
        
        # 恢复原始样式和格式
        paragraph.style = original_style
        paragraph.alignment = original_alignment
        
        # 保存文档
        if save:
            doc.save(file_path)
        
        return f"成功编辑文档 {os.path.basename(file_path)} 第 {paragraph_index+1} 段落的内容"
    except Exception as e:
        return f"编辑Word文档内容时出错: {str(e)}"

@mcp.tool()
def find_and_replace_text(
    file_path: str,
    find_text: str,
    replace_text: str,
    match_case: bool = False,
    match_whole_word: bool = False,
    save: bool = True
) -> str:
    """
    在Word文档中查找并替换文本。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        find_text: 要查找的文本
        replace_text: 替换为的文本
        match_case: 是否区分大小写，默认为False
        match_whole_word: 是否匹配整个单词，默认为False
        save: 是否保存更改，默认为True
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法在Word文档中查找替换文本，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 使用python-docx的方式（更可靠）
        doc = Document(file_path)
        replace_count = 0
        
        # 遍历所有段落和所有run
        for paragraph in doc.paragraphs:
            # 获取段落的完整文本
            full_text = paragraph.text
            if not match_case:
                search_text = find_text.lower()
                full_text_lower = full_text.lower()
            else:
                search_text = find_text
                full_text_lower = full_text
            
            # 如果段落中包含要查找的文本
            if search_text in full_text_lower:
                # 清除所有runs
                for run in paragraph.runs:
                    run.clear()
                
                # 如果需要区分大小写
                if match_case:
                    # 直接替换
                    new_text = full_text.replace(find_text, replace_text)
                    replace_count += full_text.count(find_text)
                else:
                    # 不区分大小写的替换
                    current_pos = 0
                    new_text = ""
                    while True:
                        # 在剩余文本中查找目标字符串
                        pos = full_text_lower[current_pos:].find(search_text)
                        if pos == -1:
                            # 没有找到更多匹配，添加剩余文本
                            new_text += full_text[current_pos:]
                            break
                        
                        # 添加匹配位置之前的文本
                        new_text += full_text[current_pos:current_pos + pos]
                        # 添加替换文本
                        new_text += replace_text
                        # 更新位置
                        current_pos += pos + len(find_text)
                        replace_count += 1
                
                # 添加新的run，包含替换后的文本
                paragraph.add_run(new_text)
        
        # 遍历所有表格单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 获取段落的完整文本
                        full_text = paragraph.text
                        if not match_case:
                            search_text = find_text.lower()
                            full_text_lower = full_text.lower()
                        else:
                            search_text = find_text
                            full_text_lower = full_text
                        
                        # 如果段落中包含要查找的文本
                        if search_text in full_text_lower:
                            # 清除所有runs
                            for run in paragraph.runs:
                                run.clear()
                            
                            # 如果需要区分大小写
                            if match_case:
                                # 直接替换
                                new_text = full_text.replace(find_text, replace_text)
                                replace_count += full_text.count(find_text)
                            else:
                                # 不区分大小写的替换
                                current_pos = 0
                                new_text = ""
                                while True:
                                    # 在剩余文本中查找目标字符串
                                    pos = full_text_lower[current_pos:].find(search_text)
                                    if pos == -1:
                                        # 没有找到更多匹配，添加剩余文本
                                        new_text += full_text[current_pos:]
                                        break
                                    
                                    # 添加匹配位置之前的文本
                                    new_text += full_text[current_pos:current_pos + pos]
                                    # 添加替换文本
                                    new_text += replace_text
                                    # 更新位置
                                    current_pos += pos + len(find_text)
                                    replace_count += 1
                            
                            # 添加新的run，包含替换后的文本
                            paragraph.add_run(new_text)
        
        # 保存文档
        if save:
            doc.save(file_path)
        
        return f"成功在文档 {os.path.basename(file_path)} 中替换了 {replace_count} 处文本"
    
    except Exception as e:
        return f"在Word文档中查找替换文本时出错: {str(e)}"

@mcp.tool()
def delete_paragraph(
    file_path: str,
    paragraph_index: int,
    save: bool = True
) -> str:
    """
    删除Word文档中指定的段落。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        paragraph_index: 要删除的段落索引 (从0开始计数)
        save: 是否保存更改，默认为True
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法编辑Word文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"错误: 无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落"
        
        # 获取要删除的段落
        paragraph = doc.paragraphs[paragraph_index]
        
        # 删除段落
        p = paragraph._element
        p.getparent().remove(p)
        
        # 删除对象的引用
        paragraph._p = None
        paragraph._element = None
        
        # 保存文档
        if save:
            doc.save(file_path)
        
        return f"成功从文档 {os.path.basename(file_path)} 中删除第 {paragraph_index+1} 段落"
    except Exception as e:
        return f"删除Word文档段落时出错: {str(e)}"

@mcp.tool()
def insert_table_of_contents(
    file_path: str,
    title: str = "目录",
    levels: int = 3,
    after_paragraph: int = 0
) -> str:
    """
    在Word文档中插入目录。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        title: 目录标题
        levels: 目录级别数 (1-9)
        after_paragraph: 在指定段落后插入目录，默认为文档开头第一段后
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法插入目录，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 校验参数
    if levels < 1 or levels > 9:
        return "错误: 目录级别数必须在1至9之间"
    
    try:
        # 使用python-docx库直接添加目录XML标记比较复杂
        # 尝试使用Word COM对象添加目录
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            
            # 将光标移动到指定段落后
            if after_paragraph >= 0 and after_paragraph < doc.Paragraphs.Count:
                # 修复：不再使用Select方法，改用EndOf和InsertAfter方法以保留原段落
                range_to_insert = doc.Paragraphs(after_paragraph + 1).Range
                range_to_insert.Collapse(0)  # 0表示wdCollapseEnd，折叠到段落末尾
                
                # 插入换行符创建新段落
                range_to_insert.InsertParagraphAfter()
                range_to_insert.Collapse(0)  # 再次折叠到末尾
                
                # 插入标题
                if title:
                    range_to_insert.Text = title
                    range_to_insert.InsertParagraphAfter()
                    range_to_insert.Collapse(0)
                
                # 插入目录
                toc_range = range_to_insert
                toc_range.Fields.Add(Range=toc_range, Type=-1, Text=f"TOC \\o \"1-{levels}\" \\h", PreserveFormatting=True)
            else:
                # 在文档开头插入目录
                range_to_insert = doc.Paragraphs(1).Range
                range_to_insert.Collapse(1)  # 1表示wdCollapseStart，折叠到段落开头
                
                # 插入标题
                if title:
                    range_to_insert.Text = title
                    range_to_insert.InsertParagraphAfter()
                    range_to_insert.Collapse(0)
                
                # 插入目录
                toc_range = range_to_insert
                toc_range.Fields.Add(Range=toc_range, Type=-1, Text=f"TOC \\o \"1-{levels}\" \\h", PreserveFormatting=True)
            
            # 更新目录
            if doc.TablesOfContents.Count > 0:
                doc.TablesOfContents(1).Update()
            
            # 保存并关闭
            doc.Save()
            doc.Close()
            word.Quit()
            
            pythoncom.CoUninitialize()
            
            return f"成功在文档 {os.path.basename(file_path)} 中插入目录"
        
        except ImportError:
            # 使用python-docx的方式添加目录（功能受限）
            doc = Document(file_path)
            
            # 检查指定段落是否有效
            if after_paragraph >= len(doc.paragraphs):
                return f"错误: 无效的段落索引 {after_paragraph}，文档共有 {len(doc.paragraphs)} 个段落"
            
            # 修复：在指定位置插入目录标题，确保不删除原段落
            if after_paragraph == 0:
                # 在文档开头插入
                if title:
                    # 将标题插入到第一段前
                    heading_para = doc.add_paragraph(title, style="Heading 1")
                    first_para = doc.paragraphs[1]._p  # 获取原第一段
                    heading_para._p.addnext(first_para)  # 确保原第一段仍然在标题后
            else:
                # 在指定段落后插入
                paragraph = doc.paragraphs[after_paragraph]
                if title:
                    new_para = doc.add_paragraph()
                    # 确保新段落在指定段落之后
                    paragraph._p.addnext(new_para._p)
                    new_para.text = title
                    new_para.style = "Heading 1"
            
            # 创建目录字段
            toc_para = doc.add_paragraph()
            # 确保目录段落在标题之后，不覆盖原有内容
            if title:
                if after_paragraph == 0:
                    doc.paragraphs[1]._p.addnext(toc_para._p)
                else:
                    doc.paragraphs[after_paragraph + 1]._p.addnext(toc_para._p)
            else:
                paragraph._p.addnext(toc_para._p)
            
            toc_run = toc_para.add_run()
            
            # 添加目录字段XML（这是一个简化版，功能受限）
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            toc_run._r.append(fldChar)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = f' TOC \\o "1-{levels}" \\h '
            toc_run._r.append(instrText)
            
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            toc_run._r.append(fldChar)
            
            # 保存文档
            doc.save(file_path)
            
            # 注意：使用python-docx添加的目录需要在Word中手动更新
            return f"成功在文档 {os.path.basename(file_path)} 中插入目录（需要在Word中手动更新）"
    
    except Exception as e:
        return f"插入目录时出错: {str(e)}"

@mcp.tool()
def add_header_footer(
    file_path: str,
    header_text: str = None,
    footer_text: str = None,
    page_numbers: bool = False
) -> str:
    """
    为Word文档添加页眉和页脚。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        header_text: 页眉文本（可选）
        footer_text: 页脚文本（可选）
        page_numbers: 是否在页脚添加页码
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法添加页眉页脚，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 检查是否提供了有效的参数
    if header_text is None and footer_text is None and not page_numbers:
        return "错误: 请至少提供页眉文本、页脚文本或启用页码"
    
    try:
        # 尝试使用Word COM对象添加页眉页脚（功能最完整）
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            
            # 添加页眉
            if header_text:
                for section in range(1, doc.Sections.Count + 1):
                    header = doc.Sections(section).Headers(1)  # 1表示主页眉
                    header.Range.Text = header_text
            
            # 添加页脚
            if footer_text or page_numbers:
                for section in range(1, doc.Sections.Count + 1):
                    footer = doc.Sections(section).Footers(1)  # 1表示主页脚
                    
                    if footer_text:
                        footer.Range.Text = footer_text
                    
                    # 添加页码
                    if page_numbers:
                        footer.PageNumbers.Add()
            
            # 保存并关闭
            doc.Save()
            doc.Close()
            word.Quit()
            
            pythoncom.CoUninitialize()
            
            return f"成功为文档 {os.path.basename(file_path)} 添加页眉页脚"
        
        except ImportError:
            # 使用python-docx的方式添加页眉页脚（功能受限）
            doc = Document(file_path)
            
            # 获取所有节
            sections = doc.sections
            
            # 为每个节添加页眉页脚
            for section in sections:
                # 添加页眉
                if header_text:
                    header = section.header
                    header_para = header.paragraphs[0]
                    header_para.text = header_text
                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加页脚
                if footer_text:
                    footer = section.footer
                    footer_para = footer.paragraphs[0]
                    footer_para.text = footer_text
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加页码（python-docx对页码支持有限）
                if page_numbers:
                    footer = section.footer
                    footer_para = footer.paragraphs[0] if footer_text else footer.add_paragraph()
                    
                    # 添加页码字段
                    run = footer_para.add_run()
                    
                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(qn('w:fldCharType'), 'begin')
                    run._r.append(fldChar)
                    
                    instrText = OxmlElement('w:instrText')
                    instrText.text = ' PAGE '
                    run._r.append(instrText)
                    
                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(qn('w:fldCharType'), 'end')
                    run._r.append(fldChar)
                    
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存文档
            doc.save(file_path)
            
            return f"成功为文档 {os.path.basename(file_path)} 添加页眉页脚"
    
    except Exception as e:
        return f"添加页眉页脚时出错: {str(e)}"

@mcp.tool()
def set_page_layout(
    file_path: str,
    orientation: str = None,
    page_width: float = None,
    page_height: float = None,
    left_margin: float = None,
    right_margin: float = None,
    top_margin: float = None,
    bottom_margin: float = None,
    section_index: int = 0
) -> str:
    """
    设置Word文档的页面布局。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        orientation: 页面方向，可选值: "portrait"(纵向), "landscape"(横向)
        page_width: 页面宽度（厘米，自定义纸张尺寸时使用）
        page_height: 页面高度（厘米，自定义纸张尺寸时使用）
        left_margin: 左边距（厘米）
        right_margin: 右边距（厘米）
        top_margin: 上边距（厘米）
        bottom_margin: 下边距（厘米）
        section_index: 节索引，默认为0（第一节）
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法设置页面布局，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 校验方向参数
    orientation_map = {
        "portrait": WD_ORIENTATION.PORTRAIT,
        "landscape": WD_ORIENTATION.LANDSCAPE
    }
    
    if orientation and orientation.lower() not in orientation_map:
        return f"错误: 无效的页面方向 '{orientation}'，可选值为: portrait, landscape"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 检查节索引是否有效
        if section_index < 0 or section_index >= len(doc.sections):
            return f"错误: 无效的节索引 {section_index}，文档共有 {len(doc.sections)} 节"
        
        # 获取指定的节
        section = doc.sections[section_index]
        
        # 设置页面方向
        if orientation:
            section.orientation = orientation_map[orientation.lower()]
        
        # 设置页面尺寸
        if page_width and page_height:
            section.page_width = Cm(page_width)
            section.page_height = Cm(page_height)
        
        # 设置页边距
        if left_margin is not None:
            section.left_margin = Cm(left_margin)
        
        if right_margin is not None:
            section.right_margin = Cm(right_margin)
        
        if top_margin is not None:
            section.top_margin = Cm(top_margin)
        
        if bottom_margin is not None:
            section.bottom_margin = Cm(bottom_margin)
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功设置文档 {os.path.basename(file_path)} 第 {section_index+1} 节的页面布局"
    except Exception as e:
        return f"设置页面布局时出错: {str(e)}"

@mcp.tool()
def merge_documents(
    main_file_path: str,
    files_to_merge: List[str]
) -> str:
    """
    合并多个Word文档。
    
    Args:
        main_file_path: 主文档的完整路径或相对于输出目录的路径（合并后的文档将保存为该文件）
        files_to_merge: 要合并的文档路径列表
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法合并文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(main_file_path):
        # 从环境变量获取基础路径
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        
        # 构建完整路径
        main_file_path = os.path.join(base_path, main_file_path)
    
    # 校验参数
    if not files_to_merge:
        return "错误: 请提供至少一个要合并的文档"
    
    # 处理文件路径
    processed_files = []
    for file_path in files_to_merge:
        if not os.path.isabs(file_path):
            # 从环境变量获取基础路径
            base_path = os.environ.get('OFFICE_EDIT_PATH')
            if not base_path:
                base_path = os.path.join(os.path.expanduser('~'), '桌面')
            
            # 构建完整路径
            file_path = os.path.join(base_path, file_path)
        
        # 确保文件存在
        if not os.path.exists(file_path):
            return f"错误: 文件 {file_path} 不存在"
        
        processed_files.append(file_path)
    
    try:
        # 尝试使用Word COM对象合并文档（功能最完整）
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # 检查主文档是否存在，如果不存在则创建
            if not os.path.exists(main_file_path):
                doc = word.Documents.Add()
                doc.SaveAs(main_file_path)
            else:
                doc = word.Documents.Open(main_file_path)
            
            # 记录成功合并的文档数量
            merged_count = 0
            
            # 合并每个文档
            for file_path in processed_files:
                # 将光标移动到文档末尾
                word.Selection.EndKey(Unit=6)  # 6表示wdStory，即整个文档
                
                # 插入分节符
                if merged_count > 0:
                    word.Selection.InsertBreak(Type=2)  # 2表示wdSectionBreakNextPage
                
                # 插入文档内容
                word.Selection.InsertFile(file_path)
                merged_count += 1
            
            # 保存并关闭
            doc.Save()
            doc.Close()
            word.Quit()
            
            pythoncom.CoUninitialize()
            
            return f"成功将 {merged_count} 个文档合并到 {os.path.basename(main_file_path)}"
        
        except ImportError:
            # 使用python-docx方式合并文档（功能受限）
            # 检查主文档是否存在，如果不存在则创建
            if os.path.exists(main_file_path):
                main_doc = Document(main_file_path)
            else:
                main_doc = Document()
            
            # 记录成功合并的文档数量
            merged_count = 0
            
            # 合并每个文档
            for file_path in processed_files:
                # 打开要合并的文档
                doc_to_merge = Document(file_path)
                
                # 插入分节符（如果不是第一个文档）
                if merged_count > 0:
                    main_doc.add_section()
                
                # 复制所有段落
                for paragraph in doc_to_merge.paragraphs:
                    # 创建新段落
                    new_paragraph = main_doc.add_paragraph()
                    
                    # 复制文本和格式
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name
                    
                    # 复制段落格式
                    new_paragraph.style = paragraph.style
                    new_paragraph.alignment = paragraph.alignment
                
                # 复制所有表格
                for table in doc_to_merge.tables:
                    # 创建新表格
                    new_table = main_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style
                    
                    # 复制所有单元格内容
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = cell.text
                
                merged_count += 1
            
            # 保存合并后的文档
            main_doc.save(main_file_path)
            
            return f"成功将 {merged_count} 个文档合并到 {os.path.basename(main_file_path)}"
    
    except Exception as e:
        return f"合并文档时出错: {str(e)}"
@mcp.tool()
def batch_process_document_structure(
    file_path: str,
    structure: List[Dict[str, Any]],
    clear_existing: bool = False
) -> str:
    """
    批量处理文档结构，一次性添加多种类型的内容到Word文档。
    支持标题、段落、表格、列表、图片等多种元素的批量添加。
    
    Args:
        file_path: Word文档的完整路径或相对于输出目录的路径
        structure: 文档结构数组，包含多个文档元素的配置
        clear_existing: 是否清空现有内容，默认为False（在现有内容后追加）
    
    结构元素格式示例:
    [
        {
            "type": "heading",  # 元素类型
            "content": "标题内容",
            "level": 1,  # 标题级别
            "font_size": 16,
            "font_family": "微软雅黑",
            "alignment": "center",
            "bold": true,
            "color": "#FF0000"
        },
        {
            "type": "paragraph",
            "content": "段落内容",
            "font_size": 12,
            "font_family": "宋体",
            "alignment": "left",
            "line_spacing": 1.5,
            "space_before": 6,
            "space_after": 6,
            "bold": false,
            "italic": false
        },
        {
            "type": "table",
            "rows": 3,
            "cols": 3,
            "data": [["列1", "列2", "列3"], ["数据1", "数据2", "数据3"]],
            "style": "Table Grid"
        },
        {
            "type": "list",
            "list_type": "bullet",  # bullet 或 number
            "items": ["项目1", "项目2", "项目3"]
        },
        {
            "type": "image",
            "path": "图片路径",
            "width": 10,  # 厘米
            "height": 8   # 厘米
        },
        {
            "type": "page_break"
        }
    ]
    
    Returns:
        操作结果信息
    """
    # 检查是否安装了必要的库
    if not docx_installed:
        return "错误: 无法批量处理文档，请先安装python-docx库: pip install python-docx"
    
    # 检查是否提供了完整路径
    if not os.path.isabs(file_path):
        base_path = os.environ.get('OFFICE_EDIT_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        file_path = os.path.join(base_path, file_path)
    
    # 确保文件存在
    if not os.path.exists(file_path):
        return f"错误: 文件 {file_path} 不存在"
    
    # 校验结构参数
    if not structure or not isinstance(structure, list):
        return "错误: 请提供有效的文档结构数组"
    
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 如果需要清空现有内容
        if clear_existing:
            # 删除所有段落
            for paragraph in doc.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            
            # 删除所有表格
            for table in doc.tables:
                t = table._element
                t.getparent().remove(t)
        
        # 批量处理每个结构元素
        processed_count = 0
        
        for item in structure:
            element_type = item.get('type', '').lower()
            
            if element_type == 'heading':
                _process_heading_element(doc, item)
                processed_count += 1
                
            elif element_type == 'paragraph':
                _process_paragraph_element(doc, item)
                processed_count += 1
                
            elif element_type == 'table':
                _process_table_element(doc, item)
                processed_count += 1
                
            elif element_type == 'list':
                _process_list_element(doc, item)
                processed_count += 1
                
            elif element_type == 'image':
                _process_image_element(doc, item, file_path)
                processed_count += 1
                
            elif element_type == 'page_break':
                doc.add_page_break()
                processed_count += 1
                
            else:
                # 未知类型，跳过
                continue
        
        # 保存文档
        doc.save(file_path)
        
        return f"成功批量处理文档 {os.path.basename(file_path)}，共处理 {processed_count} 个元素"
        
    except Exception as e:
        return f"批量处理文档结构时出错: {str(e)}"

def _process_heading_element(doc: Document, item: Dict[str, Any]):
    """处理标题元素"""
    content = item.get('content', '')
    level = item.get('level', 1)
    
    # 添加标题
    heading = doc.add_heading(content, level=level)
    
    # 应用格式
    _apply_text_formatting(heading, item)

def _process_paragraph_element(doc: Document, item: Dict[str, Any]):
    """处理段落元素"""
    content = item.get('content', '')
    
    # 添加段落
    paragraph = doc.add_paragraph(content)
    
    # 应用格式
    _apply_text_formatting(paragraph, item)
    _apply_paragraph_formatting(paragraph, item)

def _process_table_element(doc: Document, item: Dict[str, Any]):
    """处理表格元素"""
    rows = item.get('rows', 2)
    cols = item.get('cols', 2)
    data = item.get('data', [])
    style = item.get('style', 'Table Grid')
    
    # 添加表格
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    
    # 填充数据
    for i, row_data in enumerate(data):
        if i < rows:
            for j, cell_data in enumerate(row_data):
                if j < cols:
                    table.cell(i, j).text = str(cell_data)

def _process_list_element(doc: Document, item: Dict[str, Any]):
    """处理列表元素"""
    items = item.get('items', [])
    list_type = item.get('list_type', 'bullet')
    
    style_name = 'List Bullet' if list_type == 'bullet' else 'List Number'
    
    for list_item in items:
        paragraph = doc.add_paragraph(str(list_item), style=style_name)

def _process_image_element(doc: Document, item: Dict[str, Any], base_file_path: str):
    """处理图片元素"""
    image_path = item.get('path', '')
    width = item.get('width')
    height = item.get('height')
    
    # 处理相对路径
    if not os.path.isabs(image_path):
        base_dir = os.path.dirname(base_file_path)
        image_path = os.path.join(base_dir, image_path)
    
    # 检查图片是否存在
    if not os.path.exists(image_path):
        return
    
    # 添加图片
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    if width and height:
        run.add_picture(image_path, width=Cm(width), height=Cm(height))
    elif width:
        run.add_picture(image_path, width=Cm(width))
    elif height:
        run.add_picture(image_path, height=Cm(height))
    else:
        run.add_picture(image_path)

def _apply_text_formatting(paragraph, item: Dict[str, Any]):
    """应用文本格式"""
    font_size = item.get('font_size')
    font_family = item.get('font_family')
    bold = item.get('bold')
    italic = item.get('italic')
    underline = item.get('underline')
    color = item.get('color')
    alignment = item.get('alignment')
    
    # 应用字体格式到所有runs
    for run in paragraph.runs:
        if font_size:
            run.font.size = Pt(int(font_size))
        if font_family:
            run.font.name = font_family
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_family)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if underline is not None:
            run.font.underline = underline
        if color:
            try:
                if color.startswith("#"):
                    color = color[1:]
                r = int(color[0:2], 16)
                g = int(color[2:4], 16)
                b = int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except:
                pass
    
    # 应用段落对齐
    if alignment:
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            paragraph.alignment = alignment_map[alignment.lower()]

def _apply_paragraph_formatting(paragraph, item: Dict[str, Any]):
    """应用段落格式"""
    line_spacing = item.get('line_spacing')
    space_before = item.get('space_before')
    space_after = item.get('space_after')
    
    if line_spacing:
        paragraph.paragraph_format.line_spacing = float(line_spacing)
    if space_before:
        paragraph.paragraph_format.space_before = Pt(float(space_before))
    if space_after:
        paragraph.paragraph_format.space_after = Pt(float(space_after))


if __name__ == "__main__":
    # 运行MCP服务器
    print("启动OFFICE EDITOR服务器...")
    try:
        mcp.run()
    except TypeError as e:
        if "subscriptable" in str(e):
            # Fallback for newer anyio versions
            import asyncio
            asyncio.run(mcp.run_stdio_async())
        else:
            raise